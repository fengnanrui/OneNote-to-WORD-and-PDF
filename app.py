#!/usr/bin/env pythonw
# -*- coding: utf-8 -*-

import sys
import os
import logging
import base64
import traceback
import ctypes
import tempfile
import subprocess
import re
import html
from pathlib import Path
from typing import Dict, List, Optional
import xml.etree.ElementTree as ET

# ======= 依赖（尽量最少） =======
try:
    import win32com.client  # type: ignore
    import comtypes.client  # type: ignore
    COM_AVAILABLE = True
except Exception:
    COM_AVAILABLE = False

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QVBoxLayout, QHBoxLayout, QWidget,
    QPushButton, QLabel, QTreeWidget, QTreeWidgetItem, QProgressBar,
    QTextEdit, QGroupBox, QCheckBox, QFileDialog, QMessageBox, QSplitter,
    QTreeWidgetItemIterator
)
from PyQt5.QtCore import Qt, QTimer, QDateTime, QThread, pyqtSignal, QCoreApplication

# Word/PDF 依赖
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.section import WD_ORIENT, WD_SECTION

from reportlab.lib.pagesizes import A4
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
    Table, TableStyle, KeepInFrame, PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT


# ======= 工具函数 =======
def is_admin() -> bool:
    try:
        return ctypes.windll.shell32.IsUserAnAdmin()
    except Exception:
        return False


def check_onenote_process() -> bool:
    try:
        r = subprocess.run(['tasklist', '/FI', 'IMAGENAME eq ONENOTE.EXE'], capture_output=True, text=True)
        return 'ONENOTE.EXE' in r.stdout
    except Exception:
        return False


# ======= 一些轻量 UI 组件 =======
from PyQt5.QtCore import QTimer
from PyQt5.QtGui import QPainter, QPen, QColor


class LoadingIndicator(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedSize(26, 26)
        self.angle = 0
        self._timer = QTimer(self)
        self._timer.timeout.connect(self._tick)
        self._timer.setInterval(60)
        self.hide()

    def start(self, show_text=True):
        self._timer.start()
        self.show()

    def stop(self):
        self._timer.stop()
        self.hide()

    def _tick(self):
        self.angle = (self.angle + 10) % 360
        self.update()

    def paintEvent(self, e):
        p = QPainter(self)
        p.setRenderHint(QPainter.Antialiasing)
        pen = QPen(QColor(70, 130, 180))
        pen.setWidth(3)
        p.setPen(pen)
        r = 10
        p.translate(self.rect().center())
        p.rotate(self.angle)
        p.drawArc(-r, -r, 2*r, 2*r, 0, 120*16)


class StatusIndicator(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self._label = QLabel('', self)
        self._spinner = LoadingIndicator(self)
        layout = QHBoxLayout(self)
        layout.setContentsMargins(6, 6, 6, 6)
        layout.addWidget(self._spinner)
        layout.addWidget(self._label)
        self.hide()

    def show_loading(self, text: str):
        if self._label.text() != text:  # 只在文本变化时更新
            self._label.setText(text)
        if not self._spinner._timer.isActive():
            self._spinner.start()
        if not self.isVisible():
            self.show()
        # 不调用processEvents，避免阻塞

    def hide_loading(self):
        self._spinner.stop()
        self.hide()


# ======= OneNote API（COM优先，PowerShell回退） =======
class OneNoteAPI:
    def __init__(self):
        self.app = None
        self.temp_dir = tempfile.mkdtemp(prefix='onenote_app_')
        self.logger = logging.getLogger('OneNoteAPI')

    def initialize(self) -> bool:
        try:
            if not COM_AVAILABLE:
                raise RuntimeError('COM not available')
            admin = is_admin(); running = check_onenote_process()
            self.logger.info(f'权限: admin={admin}, running={running}')
            # 尝试三种COM
            try:
                self.app = win32com.client.gencache.EnsureDispatch('OneNote.Application')
                _ = self.app.GetHierarchy('', 1)
                return True
            except Exception as e:
                self.logger.warning(f'gencache失败: {e}')
            try:
                self.app = win32com.client.Dispatch('OneNote.Application')
                _ = self.app.GetHierarchy('', 1)
                return True
            except Exception as e:
                self.logger.warning(f'Dispatch失败: {e}')
            try:
                self.app = comtypes.client.CreateObject('OneNote.Application')
                _ = self.app.GetHierarchy('', 1)
                return True
            except Exception as e:
                self.logger.warning(f'comtypes失败: {e}')
            # 退到仅PS
            self.app = None
            return True
        except Exception as e:
            self.logger.error(f'初始化失败: {e}')
            return False

    def _ps(self, script: str) -> str:
        f = Path(self.temp_dir) / 'tmp.ps1'
        f.write_text(script, encoding='utf-8')
        r = subprocess.run(['powershell', '-ExecutionPolicy', 'Bypass', '-File', str(f)], capture_output=True, text=True, encoding='utf-8', errors='replace')
        return (r.stdout or '')

    def _get_hierarchy_ps(self, obj_id: str, scope: int) -> str:
        obj = obj_id.replace('"','""') if obj_id else ''
        sc = scope
        script = f"""
        [Console]::OutputEncoding=[System.Text.Encoding]::UTF8
        try {{
          $o=New-Object -ComObject OneNote.Application
          $x=""
          $o.GetHierarchy("{obj}",{sc},[ref]$x)
          Write-Output "SUCCESS:$x"
        }} catch {{ Write-Output "ERROR:$($_.Exception.Message)" }}
        """
        out = self._ps(script).strip()
        if out.startswith('SUCCESS:'): return out[8:]
        return ''

    def _get_page_ps(self, page_id: str) -> str:
        pid = page_id.replace('"','""')
        script = f"""
        [Console]::OutputEncoding=[System.Text.Encoding]::UTF8
        try {{
          $o=New-Object -ComObject OneNote.Application
          $x=""
          $o.GetPageContent("{pid}",[ref]$x,7)
          Write-Output "SUCCESS:$x"
        }} catch {{ Write-Output "ERROR:$($_.Exception.Message)" }}
        """
        out = self._ps(script).strip()
        if out.startswith('SUCCESS:'): return out[8:]
        return ''

    def get_notebooks(self) -> Dict:
        """获取笔记本列表，优化版本"""
        xml = ''
        try:
            if self.app:
                try:
                    # COM调用可能很慢，但在子线程中执行，不会阻塞UI
                    xml = self.app.GetHierarchy('', 4)
                except Exception:
                    pass
            if not xml:
                xml = self._get_hierarchy_ps('', 4)
        except Exception:
            xml=''
        
        if not xml:
            return {}
        
        # 解析XML - 优化版本
        try:
            root = ET.fromstring(xml)
        except Exception:
            return {}
        
        def findall_local(p, name):
            return [e for e in p.iter() if (isinstance(e.tag,str) and (e.tag.endswith('}'+name) or e.tag==name or e.tag.split('}')[-1]==name))]
        
        notebooks={}
        for nb in findall_local(root, 'Notebook'):
            nb_id = nb.get('ID')
            nb_name = nb.get('name')
            if not nb_id or not nb_name: 
                continue
            
            sections={}
            for sec in findall_local(nb,'Section'):
                sid = sec.get('ID')
                sname = sec.get('name')
                if not sid or not sname: 
                    continue
                
                pages={}
                for pg in findall_local(sec,'Page'):
                    pid = pg.get('ID')
                    pname = pg.get('name')
                    if pid and pname:
                        pages[pid] = {'id':pid, 'name':pname}
                
                sections[sid] = {'id':sid, 'name':sname, 'pages':pages}
            
            notebooks[nb_id] = {'id':nb_id, 'name':nb_name, 'sections':sections}
        
        return notebooks

    def get_page_content(self, page_id: str) -> str:
        if self.app:
            try:
                c = self.app.GetPageContent(page_id, 7)
                if c and c.strip(): return c
            except Exception:
                pass
            try:
                x=''; self.app.GetPageContent(page_id, x, 7)
                if x and x.strip(): return x
            except Exception:
                pass
        return self._get_page_ps(page_id)


# ======= 解析器（Word / PDF） =======
class OneNoteContentParser:
    def __init__(self):
        self.logger = logging.getLogger('Parser')
        self.temp_files: List[str] = []
        self._setup_chinese_fonts()
    
    def _setup_chinese_fonts(self):
        """设置中文字体支持"""
        try:
            # 尝试注册系统中文字体
            chinese_fonts = [
                ('SimSun', 'C:/Windows/Fonts/simsun.ttc'),
                ('SimHei', 'C:/Windows/Fonts/simhei.ttf'), 
                ('Microsoft YaHei', 'C:/Windows/Fonts/msyh.ttc'),
                ('PingFang SC', 'C:/Windows/Fonts/PingFang.ttc'),
            ]
            
            self.chinese_font = None
            for font_name, font_path in chinese_fonts:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        self.chinese_font = font_name
                        self.logger.info(f"成功注册中文字体: {font_name}")
                        break
                    except Exception as e:
                        self.logger.debug(f"注册字体{font_name}失败: {e}")
                        continue
            
            if not self.chinese_font:
                self.chinese_font = 'Helvetica'  # 回退到默认字体
                self.logger.warning("未找到中文字体，使用默认字体")
                
        except Exception as e:
            self.chinese_font = 'Helvetica'
            self.logger.error(f"字体设置失败: {e}")
    
    def cleanup_temp_files(self):
        """清理临时文件"""
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception as e:
                self.logger.debug(f"清理临时文件失败: {e}")
        self.temp_files.clear()

    # --- 工具：命名空间无关查找 ---
    def _findall_local(self, parent: ET.Element, local_name: str) -> List[ET.Element]:
        out=[]
        for el in parent.iter():
            tag = el.tag
            if isinstance(tag,str) and (tag.endswith('}'+local_name) or tag==local_name or tag.split('}')[-1]==local_name):
                if el is not parent or tag.split('}')[-1] != local_name:
                    out.append(el)
        return out

    # --- Word ---
    def parse_page_to_docx(self, xml: str, page_name: str, out_path: str,
                           include_images=True, include_attachments=True,
                           embed_attachments=False,
                           attachments_output_dir: Optional[Path]=None) -> bool:
        try:
            root = ET.fromstring(xml)
            doc = Document()
            doc.add_heading(page_name, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

            self._write_text_word(root, doc)
            if include_images: self._images_word(root, doc)
            if include_attachments and attachments_output_dir:
                self._attachments_word(root, doc, attachments_output_dir, embed=embed_attachments)
            self._tables_word(root, doc)

            # 如检测到超宽表，额外加一页横向节重渲
            try:
                first = self._collect_first_table_rows(root)
                if first and len(first[0])>10:
                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                    section.orientation = WD_ORIENT.LANDSCAPE
                    w,h = section.page_height, section.page_width
                    section.page_width, section.page_height = w,h
                    self._tables_word(root, doc, wide_mode=True)
            except Exception:
                pass

            doc.save(out_path)
            return True
        except Exception as e:
            self.logger.error(f'DOCX失败: {e}')
            return False

    def _write_text_word(self, root: ET.Element, doc: Document):
        """改进的文本解析，保留OneNote格式"""
        # 查找所有OE（Outline Element）元素，保持结构
        outlines = self._findall_local(root, 'OE')
        if outlines:
            for oe in outlines:
                self._process_outline_element(oe, doc)
        else:
            # 兼容旧格式
            ts = self._findall_local(root,'T')
            for t in ts:
                if t.text:
                    txt = html.unescape(t.text)
                    txt = re.sub(r'<[^>]+>','',txt)
                    if txt.strip():
                        p = doc.add_paragraph()
                        # 检查格式
                        parent = t.getparent() if hasattr(t, 'getparent') else None
                        if parent is not None:
                            run = p.add_run(txt)
                            self._apply_formatting(parent, run)
                        else:
                            p.add_run(txt)
    
    def _process_outline_element(self, oe: ET.Element, doc: Document):
        """处理OneNote的Outline元素，保留层级和格式"""
        # 获取缩进级别
        indent = 0
        list_elem = self._findall_local(oe, 'List')
        if list_elem:
            for le in list_elem:
                try:
                    indent = int(le.get('indent', '0'))
                except:
                    indent = 0
        
        # 处理文本
        ts = self._findall_local(oe, 'T')
        for t in ts:
            if t.text:
                txt = html.unescape(t.text)
                txt = re.sub(r'<[^>]+>','',txt)
                if txt.strip():
                    p = doc.add_paragraph()
                    # 应用缩进
                    if indent > 0:
                        p.paragraph_format.left_indent = Inches(indent * 0.5)
                    
                    # 检查并应用样式
                    parent = t.getparent() if hasattr(t, 'getparent') else None
                    run = p.add_run(txt)
                    if parent is not None:
                        self._apply_formatting(parent, run)
    
    def _apply_formatting(self, elem: ET.Element, run):
        """应用文本格式（粗体、斜体、下划线等）"""
        try:
            tag = elem.tag.lower() if isinstance(elem.tag, str) else ''
            # 检查粗体
            if 'bold' in tag or elem.get('bold') == 'true':
                run.bold = True
            # 检查斜体
            if 'italic' in tag or elem.get('italic') == 'true':
                run.italic = True
            # 检查下划线
            if 'underline' in tag or elem.get('underline') == 'true':
                run.underline = True
            # 检查字体大小
            size = elem.get('fontSize')
            if size:
                try:
                    run.font.size = Pt(float(size))
                except:
                    pass
        except:
            pass

    def _images_word(self, root: ET.Element, doc: Document):
        """Word图片处理，智能调整图片大小"""
        imgs = self._findall_local(root, 'Image')
        
        for im in imgs:
            # 提取图片数据
            data = None
            for attr in ('data', 'Data', 'binaryData'):
                v = im.get(attr)
                if v:
                    try: 
                        data = base64.b64decode(v)
                        break
                    except Exception: 
                        pass
                        
            if not data:
                for c in im:
                    if isinstance(c.tag, str) and ('Data' in c.tag or c.tag.endswith('Data')) and c.text:
                        try: 
                            data = base64.b64decode(c.text)
                            break
                        except Exception: 
                            pass
                            
            if not data: 
                continue
                
            # 创建临时图片文件
            fd, fp = tempfile.mkstemp(suffix='.png')
            os.close(fd)
            Path(fp).write_bytes(data)
            self.temp_files.append(fp)
            
            try:
                # 获取图片尺寸进行智能缩放
                try:
                    from PIL import Image as PILImage
                    with PILImage.open(fp) as pil_img:
                        orig_width, orig_height = pil_img.size
                        aspect_ratio = orig_height / orig_width
                except ImportError:
                    # 没有PIL时使用默认比例
                    aspect_ratio = 0.75
                    orig_width = 800
                
                # Word页面可用宽度（约6.5英寸）
                max_width = 6.5
                min_width = 3.5
                
                # 根据原始宽度智能选择显示宽度
                if orig_width <= 600:
                    # 小图片，放大到合适大小
                    display_width = max(min_width, min(max_width, max_width * 0.8))
                elif orig_width <= 1200:
                    # 中等图片，使用较大尺寸
                    display_width = max_width * 0.9
                else:
                    # 大图片，使用最大宽度
                    display_width = max_width
                
                # 如果图片很高，限制宽度以防止过高
                if aspect_ratio > 1.5:  # 高图片
                    display_width = min(display_width, max_width * 0.7)
                
                doc.add_picture(fp, width=Inches(display_width))
                doc.add_paragraph()
                
            except Exception as e:
                self.logger.warning(f"添加图片失败: {e}")
                # 回退到默认处理
                try:
                    doc.add_picture(fp, width=Inches(5))
                    doc.add_paragraph()
                except Exception:
                    pass

    def _attachments_word(self, root: ET.Element, doc: Document, out_dir: Path, embed=False):
        """处理Word附件，支持内嵌和外链两种模式"""
        files = self._findall_local(root,'InsertedFile')
        if not files:
            return
            
        if out_dir:
            out_dir.mkdir(parents=True, exist_ok=True)
        
        for a in files:
            name = a.get('pathName','attachment')
            data = self._extract_attachment(a)
            if not data: 
                continue
            
            para = doc.add_paragraph()
            run = para.add_run('📎 ')
            
            if embed:
                # 内嵌模式：尝试将附件作为OLE对象嵌入
                try:
                    # 先保存到临时文件
                    import tempfile
                    fd, temp_path = tempfile.mkstemp(suffix=Path(name).suffix)
                    os.close(fd)
                    Path(temp_path).write_bytes(data)
                    
                    # 创建嵌入式链接文本
                    run2 = para.add_run(f'[内嵌附件] {name}')
                    run2.bold = True
                    run2.font.color.rgb = RGBColor(0, 0, 255)
                    
                    # 同时保存到目录（作为备份）
                    if out_dir:
                        p = out_dir / name
                        p.write_bytes(data)
                        para.add_run(f' (已保存到: {name})')
                    
                    os.unlink(temp_path)
                except Exception as e:
                    # 如果内嵌失败，回退到外链模式
                    if out_dir:
                        p = out_dir / name
                        p.write_bytes(data)
                        run = para.add_run('附件：')
                        run.bold = True
                        self._add_hyperlink(para, p.resolve().as_uri(), name)
            else:
                # 外链模式：保存文件并创建超链接
                if out_dir:
                    p = out_dir / name
                    try:
                        p.write_bytes(data)
                        run = para.add_run('附件：')
                        run.bold = True
                        self._add_hyperlink(para, p.resolve().as_uri(), name)
                    except Exception:
                        pass

    def _add_hyperlink(self, para, url: str, text: str):
        part = para.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
        link = OxmlElement('w:hyperlink'); link.set(qn('r:id'), r_id)
        r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
        u = OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u); r.append(rPr)
        t = OxmlElement('w:t'); t.text=text; r.append(t); link.append(r)
        para._p.append(link)

    def _extract_attachment(self, elem: ET.Element) -> Optional[bytes]:
        d = elem.get('binaryData')
        if d:
            try: return base64.b64decode(d)
            except Exception: pass
        for c in elem:
            if isinstance(c.tag,str) and ('Data' in c.tag or c.tag.endswith('Data')) and c.text:
                try: return base64.b64decode(c.text)
                except Exception: pass
        return None

    def _tables_word(self, root: ET.Element, doc: Document, wide_mode: bool=False):
        """修复Word表格处理，避免重复和格式问题"""
        tables = self._findall_local(root, 'Table')
        
        for tb in tables:
            rows = self._parse_table_rows_clean(tb)
            if not rows: 
                continue
            
            # 去重：移除完全重复的行
            unique_rows = []
            seen_rows = set()
            for row in rows:
                row_key = '|'.join(row)
                if row_key not in seen_rows:
                    seen_rows.add(row_key)
                    unique_rows.append(row)
            
            if not unique_rows:
                continue
                
            max_cols = max(len(row) for row in unique_rows) if unique_rows else 1
            max_rows = len(unique_rows)
            
            # 处理宽表格
            if max_cols > 12:
                # 宽表格分两部分：正常表格 + 横向页面
                if not wide_mode:
                    # 第一部分：显示前8列
                    cols_to_show = min(8, max_cols)
                    wt = doc.add_table(rows=max_rows, cols=cols_to_show)
                    self._fill_table_data(wt, unique_rows, cols_to_show)
                    
                    # 添加提示
                    p = doc.add_paragraph()
                    p.add_run(f"注：表格共{max_cols}列，完整内容请查看横向页面").italic = True
                else:
                    # 横向页面：显示所有列
                    wt = doc.add_table(rows=max_rows, cols=max_cols)
                    self._fill_table_data(wt, unique_rows, max_cols)
            else:
                # 普通表格直接显示
                wt = doc.add_table(rows=max_rows, cols=max_cols)
                self._fill_table_data(wt, unique_rows, max_cols)
            
            doc.add_paragraph()
    
    def _fill_table_data(self, table, rows, cols_limit):
        """填充表格数据的辅助函数"""
        try:
            table.style = 'Table Grid'
            table.autofit = True
        except:
            pass
            
        for i, row in enumerate(rows):
            for j in range(min(cols_limit, len(row))):
                if i < len(table.rows) and j < len(table.rows[i].cells):
                    cell_text = row[j] if j < len(row) else ''
                    cell = table.rows[i].cells[j]
                    # 清理文本
                    clean_text = self._clean_cell_text_for_word(cell_text)
                    cell.text = clean_text
                    
                    # 设置单元格格式
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.word_wrap = True
                        paragraph.paragraph_format.keep_together = True
    
    def _clean_cell_text_for_word(self, text: str) -> str:
        """清理Word单元格文本"""
        if not text:
            return ""
        
        # 去除HTML标签和转义字符
        text = html.unescape(text)
        text = re.sub(r'<[^>]+>', '', text)
        
        # 处理换行，避免单元格内换行
        text = text.replace('\r\n', ' ').replace('\r', ' ').replace('\n', ' ')
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        return text
    
    def _parse_table_rows_clean(self, table_elem: ET.Element) -> List[List[str]]:
        """清理版表格行解析，避免重复数据"""
        rows = []
        row_elements = self._findall_local(table_elem, 'Row')
        
        for row_elem in row_elements:
            cell_elements = self._findall_local(row_elem, 'Cell')
            row_data = []
            
            for cell_elem in cell_elements:
                # 使用改进的文本提取
                cell_text = self._extract_clean_cell_text(cell_elem)
                row_data.append(cell_text)
            
            # 只添加非空且有意义的行
            if row_data and any(cell.strip() for cell in row_data):
                rows.append(row_data)
        
        return rows
    
    def _extract_clean_cell_text(self, cell_elem: ET.Element) -> str:
        """提取单元格文本，避免重复内容"""
        text_parts = []
        seen_texts = set()
        
        # 查找所有T元素
        for t_elem in self._findall_local(cell_elem, 'T'):
            if t_elem.text:
                clean_text = html.unescape(t_elem.text).strip()
                clean_text = re.sub(r'<[^>]+>', '', clean_text)
                
                if clean_text and clean_text not in seen_texts:
                    seen_texts.add(clean_text)
                    text_parts.append(clean_text)
        
        # 合并文本，用空格分隔
        result = ' '.join(text_parts)
        
        # 最终清理
        result = re.sub(r'\s+', ' ', result).strip()
        
        return result

    def _parse_table_rows(self, table_elem: ET.Element) -> List[List[str]]:
        """解析表格行，改进文本提取避免换行乱格式"""
        rows = []
        for r in self._findall_local(table_elem, 'Row'):
            row = []
            for c in self._findall_local(r, 'Cell'):
                # 更全面的文本提取
                cell_text = self._extract_all_cell_text_word(c)
                row.append(cell_text)
            if row: 
                rows.append(row)
        return rows
    
    def _extract_all_cell_text_word(self, cell_elem: ET.Element) -> str:
        """为Word表格提取单元格文本，处理换行和格式"""
        try:
            text_parts = []
            
            # 递归查找所有文本内容
            def collect_text_recursive(elem):
                if elem.text and elem.text.strip():
                    text_parts.append(elem.text.strip())
                
                for child in elem:
                    collect_text_recursive(child)
                    if child.tail and child.tail.strip():
                        text_parts.append(child.tail.strip())
            
            # 专门查找T元素（OneNote文本元素）
            for t_elem in self._findall_local(cell_elem, 'T'):
                if t_elem.text:
                    clean_text = html.unescape(t_elem.text)
                    clean_text = re.sub(r'<[^>]+>', '', clean_text)
                    clean_text = clean_text.strip()
                    if clean_text:
                        text_parts.append(clean_text)
            
            # 如果T元素没找到，用递归方法
            if not text_parts:
                collect_text_recursive(cell_elem)
            
            # 处理换行：将多个文本片段用空格连接，避免换行造成的格式问题
            full_text = ' '.join([part for part in text_parts if part])
            
            # 清理多余的空白字符
            full_text = re.sub(r'\s+', ' ', full_text)
            full_text = full_text.strip()
            
            return full_text
            
        except Exception as e:
            self.logger.debug(f"Word单元格文本提取失败: {e}")
            return ""

    # --- PDF ---
    def parse_page_to_pdf(self, xml: str, page_name: str, out_path: str,
                          include_images=True, include_attachments=True,
                          attachments_output_dir: Optional[Path]=None) -> bool:
        try:
            root = ET.fromstring(xml)
            
            # 创建自定义样式，支持中文
            styles = getSampleStyleSheet()
            
            # 标题样式
            title_style = ParagraphStyle(
                'ChineseTitle',
                parent=styles['Heading1'],
                fontSize=18,
                alignment=TA_CENTER,
                fontName=self.chinese_font,
                textColor=colors.black,
                spaceAfter=12
            )
            
            # 正文样式
            normal_style = ParagraphStyle(
                'ChineseNormal',
                parent=styles['Normal'],
                fontSize=12,
                fontName=self.chinese_font,
                textColor=colors.black,
                leftIndent=0,
                rightIndent=0,
                spaceAfter=6
            )
            
            # 创建文档，使用窄边距
            doc = SimpleDocTemplate(
                out_path, 
                pagesize=A4,
                leftMargin=1.5*cm,    # 窄边距
                rightMargin=1.5*cm,   # 窄边距
                topMargin=2*cm,
                bottomMargin=2*cm
            )
            
            story = []
            
            # 添加标题
            story.append(Paragraph(page_name, title_style))
            story.append(Spacer(1, 12))

            # 解析内容
            self._write_text_pdf_enhanced(root, story, normal_style)
            if include_images: 
                self._images_pdf_enhanced(root, story)
            if include_attachments and attachments_output_dir:
                self._attachments_pdf(root, story, normal_style, attachments_output_dir)
            self._tables_pdf_enhanced(root, story, normal_style)

            doc.build(story)
            return True
        except Exception as e:
            self.logger.error(f'PDF生成失败: {e}')
            return False
    
    def _write_text_pdf_enhanced(self, root: ET.Element, story: List, normal_style: ParagraphStyle):
        """增强版PDF文本处理，更好地支持中文和格式"""
        try:
            # 查找所有文本元素，保持层次结构
            outlines = self._findall_local(root, 'OE')
            if outlines:
                for oe in outlines:
                    self._process_outline_pdf(oe, story, normal_style)
            else:
                # 兼容模式
                text_elements = self._findall_local(root, 'T')
                for t in text_elements:
                    if t.text:
                        text = self._clean_text_for_pdf(t.text)
                        if text.strip():
                            story.append(Paragraph(text, normal_style))
                            story.append(Spacer(1, 4))
        except Exception as e:
            self.logger.error(f"PDF文本处理失败: {e}")
    
    def _process_outline_pdf(self, oe: ET.Element, story: List, base_style: ParagraphStyle):
        """处理OneNote的大纲元素到PDF"""
        try:
            # 获取缩进级别
            indent_level = 0
            list_elems = self._findall_local(oe, 'List')
            if list_elems:
                try:
                    indent_level = int(list_elems[0].get('indent', '0'))
                except:
                    indent_level = 0
            
            # 处理文本
            text_elems = self._findall_local(oe, 'T')
            for t in text_elems:
                if t.text:
                    text = self._clean_text_for_pdf(t.text)
                    if text.strip():
                        # 根据缩进创建样式
                        indent_style = ParagraphStyle(
                            f'Indent{indent_level}',
                            parent=base_style,
                            leftIndent=indent_level * 20,  # 每级缩进20点
                            bulletIndent=indent_level * 15 if indent_level > 0 else 0
                        )
                        story.append(Paragraph(text, indent_style))
                        story.append(Spacer(1, 3))
        except Exception as e:
            self.logger.debug(f"大纲处理失败: {e}")
    
    def _clean_text_for_pdf(self, text: str) -> str:
        """清理文本用于PDF显示"""
        if not text:
            return ""
        
        # HTML解码
        text = html.unescape(text)
        # 移除HTML标签
        text = re.sub(r'<[^>]+>', '', text)
        # 处理换行和空白
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        text = re.sub(r'[\t\x0b\x0c]+', ' ', text)
        # 去除首尾空白但保留内部结构
        text = text.strip()
        
        return text
    
    def _images_pdf_enhanced(self, root: ET.Element, story: List):
        """增强版图片处理，支持全屏显示"""
        try:
            imgs = self._findall_local(root, 'Image')
            
            for im in imgs:
                # 提取图片数据
                data = None
                for attr in ('data', 'Data', 'binaryData'):
                    v = im.get(attr)
                    if v:
                        try: 
                            data = base64.b64decode(v)
                            break
                        except Exception: 
                            pass
                            
                if not data:
                    for c in im:
                        if isinstance(c.tag, str) and ('Data' in c.tag or c.tag.endswith('Data')) and c.text:
                            try: 
                                data = base64.b64decode(c.text)
                                break
                            except Exception: 
                                pass
                
                if not data:
                    continue
                    
                # 创建临时图片文件
                fd, temp_img = tempfile.mkstemp(suffix='.png')
                os.close(fd)
                self.temp_files.append(temp_img)
                
                try:
                    Path(temp_img).write_bytes(data)
                    
                    # 获取图片尺寸
                    try:
                        from PIL import Image as PILImage
                        with PILImage.open(temp_img) as pil_img:
                            orig_width, orig_height = pil_img.size
                    except ImportError:
                        orig_width, orig_height = 600, 400
                    
                    # 计算合适的显示尺寸
                    page_width = A4[0] - 3*cm  # 窄边距
                    page_height = A4[1] - 4*cm
                    
                    # 智能缩放
                    scale_w = page_width / orig_width
                    scale_h = page_height / orig_height
                    scale = min(scale_w, scale_h, 1.2)  # 允许适当放大
                    
                    final_width = orig_width * scale
                    final_height = orig_height * scale
                    
                    # 确保图片至少占页面70%宽度
                    min_width = page_width * 0.7
                    if final_width < min_width:
                        scale = min_width / orig_width
                        final_width = min_width
                        final_height = orig_height * scale
                    
                    img = RLImage(temp_img, width=final_width, height=final_height)
                    story.append(Spacer(1, 8))
                    story.append(img)
                    story.append(Spacer(1, 12))
                    
                except Exception as e:
                    self.logger.warning(f"处理图片失败: {e}")
                    # 使用固定大小作为回退
                    try:
                        img = RLImage(temp_img, width=5*inch, height=4*inch)
                        story.append(img)
                        story.append(Spacer(1, 12))
                    except Exception:
                        pass
                        
        except Exception as e:
            self.logger.error(f"PDF图片处理失败: {e}")
    
    def _attachments_pdf(self, root: ET.Element, story: List, normal_style: ParagraphStyle, out_dir: Path):
        """PDF附件处理：保存到目录并在文档中添加引用"""
        try:
            files = self._findall_local(root, 'InsertedFile')
            if not files:
                return
                
            out_dir.mkdir(parents=True, exist_ok=True)
            
            # 添加附件标题
            heading_style = ParagraphStyle(
                'AttachmentHeading',
                parent=normal_style,
                fontSize=14,
                textColor=colors.black,
                spaceAfter=8,
                fontName=normal_style.fontName
            )
            
            story.append(Spacer(1, 12))
            story.append(Paragraph('📎 附件列表', heading_style))
            story.append(Spacer(1, 6))
            
            for a in files:
                name = a.get('pathName', 'attachment')
                data = self._extract_attachment(a)
                if not data:
                    continue
                
                try:
                    # 保存附件到目录
                    p = out_dir / name
                    p.write_bytes(data)
                    
                    # 在PDF中添加附件信息
                    info = f"• {name} (已保存到附件目录)"
                    story.append(Paragraph(info, normal_style))
                    story.append(Spacer(1, 3))
                except Exception as e:
                    self.logger.debug(f"保存PDF附件失败: {e}")
        except Exception as e:
            self.logger.error(f"附件处理失败: {e}")
            
    def _tables_pdf_enhanced(self, root: ET.Element, story: List, normal_style: ParagraphStyle):
        """增强版PDF表格处理，完整保留数据并支持中文"""
        try:
            table_elements = self._findall_local(root, 'Table')
            if not table_elements:
                return
            
            # 创建表格专用样式 - 优化字体和行距
            cell_style = ParagraphStyle(
                'TableCell',
                parent=normal_style,
                fontSize=8,  # 稍小字体，节省空间
                leading=10,  # 紧凑行距
                fontName=normal_style.fontName,
                leftIndent=1,
                rightIndent=1,
                spaceAfter=1,
                spaceBefore=1,
                wordWrap=True  # 启用自动换行
            )
            
            header_style = ParagraphStyle(
                'TableHeader',
                parent=cell_style,
                fontSize=9,  # 标题稍大一点
                leading=11,
                textColor=colors.white,
                alignment=TA_CENTER,
                fontWeight='bold'
            )
            
            for table_idx, table_elem in enumerate(table_elements):
                # 使用和Word相同的清理方法
                rows_data = self._parse_table_rows_clean(table_elem)
                if not rows_data:
                    continue
                
                # 去重处理，避免重复行
                unique_rows = []
                seen_rows = set()
                for row in rows_data:
                    row_key = '|'.join(row)
                    if row_key not in seen_rows:
                        seen_rows.add(row_key)
                        unique_rows.append(row)
                
                if not unique_rows:
                    continue
                
                # 数据预处理
                cleaned_rows = []
                max_cols = 0
                
                for row in unique_rows:
                    cleaned_row = []
                    for cell in row:
                        # 清理文本用于PDF显示
                        cleaned_text = self._clean_cell_text_for_pdf(cell)
                        cleaned_row.append(cleaned_text)
                    cleaned_rows.append(cleaned_row)
                    max_cols = max(max_cols, len(cleaned_row))
                
                # 补齐所有行到相同列数
                for row in cleaned_rows:
                    while len(row) < max_cols:
                        row.append('')
                
                # 大幅简化分段策略，减少PDF页面混乱
                if max_cols <= 8:
                    MAX_COLS_PER_SEGMENT = max_cols  # 8列以内不分段
                elif max_cols <= 12:
                    MAX_COLS_PER_SEGMENT = 8  # 12列以内分2段
                else:
                    MAX_COLS_PER_SEGMENT = 6  # 超宽表格每段6列
                    
                MAX_ROWS = 120  # 增加行数，减少表格数量
                
                # 限制行数
                if len(cleaned_rows) > MAX_ROWS:
                    cleaned_rows = cleaned_rows[:MAX_ROWS]
                    truncated = True
                else:
                    truncated = False
                
                # 计算分段
                col_segments = list(range(0, max_cols, MAX_COLS_PER_SEGMENT))
                
                for seg_idx, start_col in enumerate(col_segments):
                    end_col = min(start_col + MAX_COLS_PER_SEGMENT, max_cols)
                    
                    # 提取当前段的数据
                    segment_data = []
                    for row in cleaned_rows:
                        segment_row = row[start_col:end_col]
                        segment_data.append(segment_row)
                    
                    if not segment_data:
                        continue
                    
                    # 转换为Paragraph对象
                    table_flow = []
                    for row_idx, row in enumerate(segment_data):
                        flow_row = []
                        for col_idx, cell_text in enumerate(row):
                            if row_idx == 0:  # 假设第一行是标题
                                para = Paragraph(cell_text or ' ', header_style)
                            else:
                                para = Paragraph(cell_text or ' ', cell_style)
                            flow_row.append(para)
                        table_flow.append(flow_row)
                    
                    if not table_flow:
                        continue
                    
                    # 简化段标题（只在必要时显示）
                    if len(col_segments) > 1 and max_cols > 8:
                        if seg_idx == 0:
                            # 只在第一段显示表格标题
                            seg_title = f"表格 {table_idx + 1} (共{max_cols}列，分{len(col_segments)}部分显示)"
                        else:
                            # 后续段只显示列范围
                            seg_title = f"续表 (列 {start_col + 1}-{end_col})"
                        
                        title_style = ParagraphStyle(
                            'SegmentTitle',
                            parent=normal_style,
                            fontSize=10,
                            textColor=colors.darkblue,
                            spaceAfter=3,
                            spaceBefore=6
                        )
                        story.append(Paragraph(seg_title, title_style))
                    
                    try:
                        # 动态计算列宽
                        available_width = A4[0] - 3*cm
                        col_width = available_width / len(table_flow[0])
                        col_widths = [col_width] * len(table_flow[0])
                        
                        # 创建表格
                        pdf_table = Table(table_flow, colWidths=col_widths, repeatRows=1)
                        
                        # 简化表格样式，提高可读性
                        table_style = TableStyle([
                            # 外边框
                            ('BOX', (0, 0), (-1, -1), 1, colors.black),
                            # 内部网格线
                            ('INNERGRID', (0, 0), (-1, -1), 0.3, colors.grey),
                            # 垂直对齐
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            
                            # 标题行样式
                            ('BACKGROUND', (0, 0), (-1, 0), colors.darkgrey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                            ('FONTSIZE', (0, 0), (-1, 0), 9),
                            ('FONTNAME', (0, 0), (-1, 0), normal_style.fontName),
                            
                            # 数据行样式
                            ('FONTSIZE', (0, 1), (-1, -1), 8),
                            ('FONTNAME', (0, 1), (-1, -1), normal_style.fontName),
                            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                            
                            # 合理的内边距
                            ('LEFTPADDING', (0, 0), (-1, -1), 4),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                            ('TOPPADDING', (0, 0), (-1, -1), 3),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                        ])
                        
                        pdf_table.setStyle(table_style)
                        
                        # 使用KeepInFrame确保表格适应页面
                        max_height = A4[1] - 6*cm  # 留出更多空间给页边距
                        kif = KeepInFrame(available_width, max_height, [pdf_table], mode='shrink')
                        story.append(kif)
                        story.append(Spacer(1, 12))
                        
                    except Exception as render_err:
                        self.logger.warning(f"表格渲染失败，使用文本模式: {render_err}")
                        # 退化为纯文本显示
                        for row in segment_data:
                            line = ' | '.join(row)
                            story.append(Paragraph(line, normal_style))
                        story.append(Spacer(1, 8))
                
                # 如果表格被截断，添加提示
                if truncated:
                    note_style = ParagraphStyle(
                        'TruncateNote',
                        parent=normal_style,
                        fontSize=8,
                        textColor=colors.red,
                        fontStyle='italic',
                        alignment=TA_RIGHT
                    )
                    story.append(Spacer(1, 3))
                    story.append(Paragraph(f'注：表格内容过多，已显示前{MAX_ROWS}行', note_style))
                    story.append(Spacer(1, 8))
                    
        except Exception as e:
            self.logger.error(f"增强表格处理失败: {e}")

    
    def _extract_all_cell_text(self, cell_elem: ET.Element) -> str:
        """提取单元格内的所有文本，包括嵌套元素"""
        try:
            text_parts = []
            
            # 递归查找所有文本元素
            def collect_text(elem):
                if elem.text:
                    text_parts.append(elem.text.strip())
                
                # 查找所有T元素（文本元素）
                for t_elem in elem.findall('.//'):
                    if t_elem.tag.endswith('T') or t_elem.tag == 'T':
                        if t_elem.text:
                            text_parts.append(t_elem.text.strip())
                
                # 也检查元素的tail文本
                if elem.tail:
                    text_parts.append(elem.tail.strip())
            
            collect_text(cell_elem)
            
            # 合并所有文本部分
            full_text = ' '.join([part for part in text_parts if part])
            
            # 清理HTML和特殊字符
            full_text = html.unescape(full_text)
            full_text = re.sub(r'<[^>]+>', '', full_text)
            
            return full_text.strip()
            
        except Exception as e:
            self.logger.debug(f"提取单元格文本失败: {e}")
            return ""
    
    def _clean_cell_text_for_pdf(self, text: str) -> str:
        """清理单元格文本用于PDF显示，处理换行乱格式"""
        if not text:
            return " "
        
        # 基础清理
        text = html.unescape(text)
        text = re.sub(r'<[^>]+>', '', text)
        
        # 智能处理换行：
        # 1. 先将所有换行符替换为特殊标记
        text = text.replace('\r\n', '<<<LINEBREAK>>>').replace('\r', '<<<LINEBREAK>>>').replace('\n', '<<<LINEBREAK>>>')
        
        # 2. 处理多余的空白，但保留段落分隔
        text = re.sub(r'\s+', ' ', text)
        
        # 3. 恢复重要的换行为空格，避免单元格内换行乱格式
        text = text.replace('<<<LINEBREAK>>>', ' ')
        
        # 4. 最终清理
        text = re.sub(r'\s+', ' ', text)  # 压缩连续空格
        text = text.strip()
        
        # 处理特殊字符，确保PDF兼容性
        text = text.replace('\u2022', '•')  # 项目符号
        text = text.replace('\u2013', '-')  # en-dash
        text = text.replace('\u2014', '—')  # em-dash
        text = text.replace('\u201c', '"').replace('\u201d', '"')  # 引号
        text = text.replace('\u2018', "'").replace('\u2019', "'")  # 单引号
        
        # 限制长度，但保留更多内容
        MAX_CELL_LENGTH = 300  # 增加限制长度
        if len(text) > MAX_CELL_LENGTH:
            # 在空格处截断，避免截断单词
            truncate_pos = text.rfind(' ', 0, MAX_CELL_LENGTH - 3)
            if truncate_pos > MAX_CELL_LENGTH * 0.7:  # 如果找到了合适的截断位置
                text = text[:truncate_pos] + '...'
            else:
                text = text[:MAX_CELL_LENGTH - 3] + '...'
        
        return text or " "  # 确保不返回空字符串

    def _write_text_pdf(self, root: ET.Element, story: List, styles):
        ts = self._findall_local(root,'T')
        buf=[]
        for t in ts:
            if t.text:
                s = html.unescape(t.text); s=re.sub(r'<[^>]+>','',s)
                buf.append(s)
        for line in '\n'.join(buf).split('\n'):
            if line.strip():
                story.append(Paragraph(line.strip(), styles['Normal']))
                story.append(Spacer(1, 6))

    def _images_pdf(self, root: ET.Element, story: List):
        imgs = self._findall_local(root,'Image')
        for im in imgs:
            data=None
            for attr in ('data','Data','binaryData'):
                v=im.get(attr)
                if v:
                    try: data=base64.b64decode(v); break
                    except Exception: pass
            if not data:
                for c in im:
                    if isinstance(c.tag,str) and ('Data' in c.tag or c.tag.endswith('Data')) and c.text:
                        try: data=base64.b64decode(c.text); break
                        except Exception: pass
            if not data: continue
            fd, fp = tempfile.mkstemp(suffix='.png'); os.close(fd)
            Path(fp).write_bytes(data)
            img = RLImage(fp, width=4*inch, height=3*inch)
            story.append(img); story.append(Spacer(1, 12))

    def _tables_pdf(self, root: ET.Element, story: List, styles):
        """改进的PDF表格处理，确保单页显示"""
        tbls = self._findall_local(root,'Table')
        if not tbls: return
        
        # 创建紧凑的单元格样式
        cell_style = ParagraphStyle('Cell', parent=styles['Normal'], 
                                   fontSize=7, leading=8, 
                                   leftIndent=0, rightIndent=0)
        
        def clean(s:str)->str:
            s=html.unescape(s); s=re.sub(r'<[^>]+>','',s); 
            return s.strip()
        
        for tb in tbls:
            rows = self._parse_table_rows(tb)
            if not rows: continue
            
            # 预处理数据
            MAX_CELL_CHARS = 200  # 减少字符限制，使表格更紧凑
            rows = [[clean(c)[:MAX_CELL_CHARS] for c in row] for row in rows]
            total_cols = max(len(r) for r in rows) if rows else 0
            
            # 智能分段：根据列数决定每段显示多少列
            if total_cols <= 6:
                MAX_COLS = 6
            elif total_cols <= 10:
                MAX_COLS = 5
            else:
                MAX_COLS = 4
            
            MAX_ROWS = 50  # 限制行数，确保能在一页显示
            
            # 如果表格太大，进行截断
            if len(rows) > MAX_ROWS:
                rows = rows[:MAX_ROWS]
                truncated = True
            else:
                truncated = False
            
            # 分段显示
            starts = list(range(0, total_cols, MAX_COLS))
            
            for idx, st in enumerate(starts):
                ed = min(st + MAX_COLS, total_cols)
                
                # 准备段数据
                seg_data = []
                for row in rows:
                    seg_row = row[st:ed] if len(row) > st else []
                    # 补齐空单元格
                    seg_row += [''] * (ed - st - len(seg_row))
                    seg_data.append(seg_row)
                
                # 转换为Paragraph对象
                seg_flow = [[Paragraph(cell or ' ', cell_style) for cell in row] 
                           for row in seg_data]
                
                # 创建表格
                try:
                    # 动态计算列宽
                    available_width = 450  # A4页面可用宽度
                    col_width = available_width / (ed - st)
                    col_widths = [col_width] * (ed - st)
                    
                    t = Table(seg_flow, colWidths=col_widths, repeatRows=1)
                    t.setStyle(TableStyle([
                        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                        ('VALIGN', (0,0), (-1,-1), 'TOP'),
                        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                        ('FONTSIZE', (0,0), (-1,-1), 7),
                        ('LEFTPADDING', (0,0), (-1,-1), 2),
                        ('RIGHTPADDING', (0,0), (-1,-1), 2),
                        ('TOPPADDING', (0,0), (-1,-1), 1),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                        # 首行加粗（如果有标题行）
                        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                        ('FONTSIZE', (0,0), (-1,0), 8),
                        ('BOLD', (0,0), (-1,0), True),
                    ]))
                    
                    # 添加段标题
                    if len(starts) > 1:
                        story.append(Paragraph(f'表格部分 {idx+1}/{len(starts)} (列 {st+1}-{ed})', 
                                             styles['Italic']))
                        story.append(Spacer(1, 4))
                    
                    # 使用KeepInFrame确保表格在一页内
                    kif = KeepInFrame(available_width, 650, [t], mode='shrink')
                    story.append(kif)
                    story.append(Spacer(1, 12))
                    
                except Exception as e:
                    # 如果表格渲染失败，退化为文本
                    self.logger.debug(f'PDF表格渲染失败: {e}')
                    for r in seg_data:
                        story.append(Paragraph(' | '.join(r), styles['Normal']))
                    story.append(Spacer(1, 6))
            
            if truncated:
                story.append(Paragraph('... 表格内容过多，已截断显示', styles['Italic']))
                story.append(Spacer(1, 8))


# ======= GUI =======
class ModernOneNoteGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.onenote = OneNoteAPI()
        self.parser = OneNoteContentParser()
        self.selected_items=[]; self.output_dir=''
        self._busy=False
        self._loading_thread = None  # 保持线程引用
        self._populate_thread = None
        self._convert_thread = None
        self._setup_logging(); self._init_ui(); self._apply_styles()
        
        # 设置窗口属性以提升性能
        self.setAttribute(Qt.WA_OpaquePaintEvent, True)
        self.setAttribute(Qt.WA_NoSystemBackground, True)
        
        # 延迟自动检测，让UI先完全显示
        QTimer.singleShot(500, self._auto_detect)

    def _setup_logging(self):
        logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

    def _init_ui(self):
        self.setWindowTitle('WATS-OneNote_Exporter')
        self.setMinimumSize(1000, 620)
        
        # 创建主布局
        cw = QWidget()
        self.setCentralWidget(cw)
        main = QHBoxLayout(cw)
        
        # 左侧面板
        left = QWidget()
        lv = QVBoxLayout(left)
        
        title = QLabel('📚 OneNote 笔记本')
        title.setObjectName('title_label')
        lv.addWidget(title)
        
        self.refresh_btn = QPushButton('🔄 刷新笔记本')
        self.refresh_btn.clicked.connect(self._refresh)
        lv.addWidget(self.refresh_btn)
        
        self.refresh_status = StatusIndicator(left)
        self.refresh_status.setObjectName('status_bar')
        lv.addWidget(self.refresh_status)
        
        # 树控件优化设置
        self.tree = QTreeWidget()
        self.tree.setObjectName('notebook_tree')
        self.tree.setHeaderLabels(['笔记本/分区/页面','状态'])
        self.tree.itemChanged.connect(self._on_item_changed)
        
        # 设置列宽50-50分割 - 真正的50-50
        header = self.tree.header()
        header.setStretchLastSection(False)
        # 延迟设置真正的50-50分割比例
        QTimer.singleShot(100, self._setup_tree_columns)
        
        # 性能优化设置，保留展开图标
        self.tree.setUniformRowHeights(True)  # 统一行高提升性能
        self.tree.setAlternatingRowColors(False)  # 禁用交替行色，避免黑色
        self.tree.setAnimated(False)  # 禁用动画以提升性能
        self.tree.setExpandsOnDoubleClick(True)   # 允许双击展开
        self.tree.setItemsExpandable(True)        # 允许展开
        self.tree.setRootIsDecorated(True)        # 显示根装饰（展开图标）
        self.tree.setIndentation(20)              # 设置合适的缩进，显示层级
        
        lv.addWidget(self.tree)
        
        # 选择按钮
        sel_bar = QWidget()
        hb = QHBoxLayout(sel_bar)
        self.btn_all = QPushButton('✅ 全选')
        self.btn_all.clicked.connect(self._select_all)
        self.btn_none = QPushButton('❌ 取消全选')
        self.btn_none.clicked.connect(self._select_none)
        hb.addWidget(self.btn_all)
        hb.addWidget(self.btn_none)
        lv.addWidget(sel_bar)

        # 右侧面板
        right = QWidget()
        rv = QVBoxLayout(right)
        
        # 输出设置
        out_g = QGroupBox('📁 输出设置')
        out_g.setObjectName('group')
        og = QHBoxLayout(out_g)
        self.lbl_out = QLabel('未选择输出目录')
        self.btn_dir = QPushButton('选择目录')
        self.btn_dir.clicked.connect(self._choose_dir)
        og.addWidget(self.lbl_out)
        og.addWidget(self.btn_dir)
        rv.addWidget(out_g)
        
        # 导出格式
        fmt_g = QGroupBox('📄 导出格式')
        fmt_g.setObjectName('group')
        fg = QVBoxLayout(fmt_g)
        self.cb_pdf = QCheckBox('导出PDF')
        self.cb_pdf.setChecked(True)
        self.cb_docx = QCheckBox('导出Word')
        self.cb_docx.setChecked(True)
        fg.addWidget(self.cb_pdf)
        fg.addWidget(self.cb_docx)
        rv.addWidget(fmt_g)
        
        # 选项
        opt_g = QGroupBox('⚙️ 选项')
        opt_g.setObjectName('group')
        og2 = QVBoxLayout(opt_g)
        self.cb_img = QCheckBox('包含图片')
        self.cb_img.setChecked(True)
        self.cb_att = QCheckBox('包含附件')
        self.cb_att.setChecked(True)
        og2.addWidget(self.cb_img)
        og2.addWidget(self.cb_att)
        rv.addWidget(opt_g)
        
        # 转换按钮
        self.convert_btn = QPushButton('🚀 开始转换')
        self.convert_btn.clicked.connect(self._convert)
        self.convert_btn.setEnabled(False)
        rv.addWidget(self.convert_btn)
        
        # 状态和进度
        self.conv_status = StatusIndicator(right)
        self.conv_status.setObjectName('status_bar')
        rv.addWidget(self.conv_status)
        
        self.progress = QProgressBar()
        self.progress.setVisible(False)
        rv.addWidget(self.progress)
        
        # 日志
        self.log = QTextEdit()
        self.log.setObjectName('log')
        self.log.setReadOnly(True)
        self.log.setMaximumHeight(240)
        rv.addWidget(self.log)

        # 分割器 - 真正的一半一半分割
        spl = QSplitter(Qt.Horizontal)
        spl.addWidget(left)
        spl.addWidget(right)
        
        # 恢复原有布局比例
        spl.setStretchFactor(0, 1)  
        spl.setStretchFactor(1, 1)
        spl.setSizes([680, 420])  # 恢复原来的比例
        
        # 设置分割器样式
        spl.setHandleWidth(3)
        spl.setStyleSheet("""
            QSplitter::handle {
                background: #cbd5e1;
                border: 1px solid #94a3b8;
            }
            QSplitter::handle:horizontal {
                width: 3px;
            }
        """)
        
        main.addWidget(spl)

    def _setup_tree_columns(self):
        """设置树控件列为真正的50-50分割"""
        try:
            tree_width = self.tree.width() - 20  # 减去滚动条和边距
            col_width = tree_width // 2  # 每列占一半
            
            self.tree.setColumnWidth(0, col_width)
            self.tree.setColumnWidth(1, col_width)
            
            # 设置列为等比例拉伸
            header = self.tree.header()
            header.setSectionResizeMode(0, header.Stretch)
            header.setSectionResizeMode(1, header.Stretch)
            
            # 确保表头也是50-50分割
            header.setDefaultSectionSize(col_width)
            
        except Exception as e:
            self.logger.debug(f"设置列宽失败: {e}")

    def _apply_styles(self):
        """简洁的白色主题样式，完全无黑色"""
        style = """
        /* 主窗口 - 纯白背景 */
        QMainWindow { 
            background: white;
            color: #374151;
        }
        
        /* 所有控件默认白色背景 */
        QWidget {
            background: white;
            color: #374151;
        }
        
        /* 分组框 */
        QGroupBox { 
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 8px; 
            padding: 15px; 
            margin-top: 12px;
            font-size: 14px;
        }
        
        QGroupBox::title { 
            subcontrol-origin: margin; 
            left: 12px; 
            padding: 0 8px; 
            color: #1f2937; 
            font-weight: 600;
            background: white;
        }
        
        /* 标题 */
        QLabel#title_label { 
            font-size: 20px; 
            font-weight: 700; 
            color: #1f2937; 
            background: white;
            padding: 10px;
            border: 1px solid #e5e7eb;
            border-radius: 6px;
        }
        
        QLabel { 
            color: #374151;
            background: white;
        }
        
        /* 按钮 */
        QPushButton { 
            background: #3b82f6;
            color: white; 
            border: none;
            padding: 10px 16px; 
            border-radius: 6px; 
            font-weight: 600;
            font-size: 13px;
        }
        
        QPushButton:hover { 
            background: #2563eb;
        }
        
        QPushButton:disabled { 
            background: #e5e7eb; 
            color: #9ca3af;
        }
        
        /* 树控件 - 简洁样式，保留默认展开图标 */
        QTreeWidget { 
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 6px;
            font-size: 13px;
            color: #374151;
            outline: none;
        }
        
        QHeaderView::section {
            background: #f9fafb;
            color: #374151;
            border: none;
            border-right: 1px solid #e5e7eb;
            padding: 8px;
            font-weight: 600;
        }
        
        QTreeWidget::item {
            background: white;
            color: #374151;
            padding: 6px;
            height: 26px;
        }
        
        QTreeWidget::item:hover {
            background: #f3f4f6;
        }
        
        QTreeWidget::item:selected {
            background: #dbeafe;
            color: #1e40af;
        }
        
        /* 复选框 */
        QCheckBox {
            color: #374151;
            background: white;
        }
        
        QCheckBox::indicator {
            width: 16px;
            height: 16px;
            border: 1px solid #d1d5db;
            border-radius: 3px;
            background: white;
        }
        
        QCheckBox::indicator:checked {
            background: #3b82f6;
            border: 1px solid #3b82f6;
        }
        
        /* 日志区域 - 白色背景 */
        QTextEdit#log { 
            background: white;
            color: #374151; 
            border: 1px solid #e5e7eb;
            border-radius: 6px;
            font-family: 'Consolas', monospace;
            font-size: 11px;
            padding: 8px;
        }
        
        /* 进度条 */
        QProgressBar { 
            border: 1px solid #e5e7eb;
            border-radius: 6px; 
            height: 22px; 
            text-align: center;
            background: white;
            color: #374151;
        }
        
        QProgressBar::chunk { 
            background: #10b981;
            border-radius: 4px;
            margin: 1px;
        }
        """
        self.setStyleSheet(style)

    # ---- 动作 ----
    def _auto_detect(self):
        """自动检测OneNote，延迟启动避免界面卡顿"""
        # 确保UI完全初始化后再开始检测
        if self.isVisible():
            QTimer.singleShot(200, self._refresh)
        else:
            QTimer.singleShot(1000, self._auto_detect)

    def _set_busy(self, busy: bool):
        """设置忙碌状态，禁用/启用控件"""
        self._busy=busy
        
        # 禁用/启用所有控件
        controls = [self.refresh_btn,self.btn_all,self.btn_none,self.btn_dir,
                   self.cb_pdf,self.cb_docx,self.cb_img,self.cb_att,self.convert_btn]
        
        for w in controls:
            w.setEnabled(not busy)
        
        # 树控件单独处理
        self.tree.setEnabled(not busy)
        
        # 设置光标
        if busy:
            QApplication.setOverrideCursor(Qt.WaitCursor)
        else:
            QApplication.restoreOverrideCursor()
        
        # 不要在这里调用 processEvents，会导致事件循环嵌套

    def _log(self, msg: str):
        ts = QDateTime.currentDateTime().toString('hh:mm:ss')
        self.log.append(f'[{ts}] {msg}')
        self.log.verticalScrollBar().setValue(self.log.verticalScrollBar().maximum())

    def _refresh(self):
        if self._busy: return
        self._set_busy(True)
        self.refresh_status.show_loading('🔍 正在检测OneNote...')
        self.tree.clear()
        self._log('开始加载笔记本...')
        
        # 清理之前的缓存
        if hasattr(self, '_item_cache'):
            self._item_cache.clear()
        
        # 确保之前的线程已停止
        if self._loading_thread and self._loading_thread.isRunning():
            self._loading_thread.terminate()
            self._loading_thread.wait(100)
        
        self._loading_thread = _DetectWorker(self.onenote)
        self._loading_thread.progress.connect(self._on_detect_progress, Qt.QueuedConnection)
        self._loading_thread.done.connect(self._on_loaded, Qt.QueuedConnection)
        self._loading_thread.err.connect(self._on_load_err, Qt.QueuedConnection)
        self._loading_thread.start(QThread.HighPriority)  # 高优先级
    
    def _on_detect_progress(self, msg: str):
        """处理检测进度"""
        self.refresh_status.show_loading(msg)
        self._log(msg)
        # 不要调用processEvents，让Qt自动处理

    def _on_loaded(self, notebooks: dict):
        self.refresh_status.show_loading('📚 读取笔记本 0%')
        self._log('📚 开始读取笔记本结构...')
        
        # 确保之前的线程已停止
        if self._populate_thread and self._populate_thread.isRunning():
            self._populate_thread.terminate()
            self._populate_thread.wait(100)
        
        self._populate_thread = _PopulateWorker(notebooks)
        self._populate_thread.all_data.connect(self._build_tree_fast, Qt.QueuedConnection)
        self._populate_thread.progress.connect(self._on_populate_progress, Qt.QueuedConnection)
        self._populate_thread.msg.connect(self._log, Qt.QueuedConnection)
        self._populate_thread.done.connect(self._on_pop_done, Qt.QueuedConnection)
        self._populate_thread.err.connect(self._on_pop_err, Qt.QueuedConnection)
        self._populate_thread.start(QThread.HighPriority)
    
    def _on_populate_progress(self, percent: int):
        """处理构建进度"""
        if percent < 100:
            self.refresh_status.show_loading(f'📚 读取笔记本 {percent}%')
        # 不调用processEvents

    def _on_load_err(self, msg: str):
        self.refresh_status.hide_loading(); self._set_busy(False)
        self._log(f'❌ 加载失败: {msg}')

    def _build_tree_fast(self, notebooks: dict):
        """超高速构建整个树形结构"""
        try:
            # 彻底禁用所有更新和信号
            self.tree.setUpdatesEnabled(False)
            self.tree.blockSignals(True)
            self.tree.setVisible(False)  # 隐藏控件加速构建
            
            # 使用QTimer来分段处理，避免阻塞UI
            self._notebooks_data = notebooks
            self._build_items = []
            self._build_index = 0
            
            # 预处理所有项目
            for nb_id, nb_data in notebooks.items():
                nb_name = nb_data['name']
                self._build_items.append(('notebook', None, nb_id, nb_name))
                
                for sec_id, sec_data in nb_data.get('sections', {}).items():
                    sec_name = sec_data['name']
                    self._build_items.append(('section', nb_id, sec_id, sec_name))
                    
                    for page_id, page_data in sec_data.get('pages', {}).items():
                        page_name = page_data['name']
                        self._build_items.append(('page', sec_id, page_id, page_name))
            
            # 使用定时器分批处理，避免UI阻塞
            self._item_cache = {}
            self._build_timer = QTimer()
            self._build_timer.timeout.connect(self._build_batch)
            self._build_timer.start(1)  # 每1ms处理一批
            
        except Exception as e:
            self._log(f'❌ 快速构建失败: {e}')
            self._finish_build()
    
    def _build_batch(self):
        """分批构建树项目"""
        try:
            batch_size = 100  # 每次处理100个项目
            end_index = min(self._build_index + batch_size, len(self._build_items))
            
            for i in range(self._build_index, end_index):
                item_type, parent_id, item_id, item_name = self._build_items[i]
                
                if item_type == 'notebook':
                    it = QTreeWidgetItem(self.tree)
                    it.setText(0, f'📚 {item_name}')
                    it.setText(1, '笔记本')
                    it.setFlags(it.flags() | Qt.ItemIsUserCheckable)
                    it.setCheckState(0, Qt.Unchecked)
                    it.setData(0, Qt.UserRole, {'type': 'notebook', 'id': item_id, 'name': item_name})
                    it.setExpanded(True)
                    self._item_cache[item_id] = it
                    
                elif item_type == 'section':
                    parent = self._item_cache.get(parent_id)
                    if parent:
                        it = QTreeWidgetItem(parent)
                        it.setText(0, f'📁 {item_name}')
                        it.setText(1, '分区')
                        it.setFlags(it.flags() | Qt.ItemIsUserCheckable)
                        it.setCheckState(0, Qt.Unchecked)
                        it.setData(0, Qt.UserRole, {'type': 'section', 'id': item_id, 'name': item_name})
                        self._item_cache[item_id] = it
                
                elif item_type == 'page':
                    parent = self._item_cache.get(parent_id)
                    if parent:
                        it = QTreeWidgetItem(parent)
                        it.setText(0, f'📄 {item_name}')
                        it.setText(1, '页面')
                        it.setFlags(it.flags() | Qt.ItemIsUserCheckable)
                        it.setCheckState(0, Qt.Unchecked)
                        it.setData(0, Qt.UserRole, {'type': 'page', 'id': item_id, 'name': item_name})
            
            self._build_index = end_index
            
            # 更新进度
            progress = int(self._build_index * 100 / len(self._build_items))
            if progress % 10 == 0 and progress < 100:  # 每10%更新一次，但不显示100%
                self.refresh_status.show_loading(f'📚 读取笔记本 {progress}%')
            
            # 检查是否完成
            if self._build_index >= len(self._build_items):
                self._build_timer.stop()
                # 完成时不显示任何状态，直接隐藏
                self._finish_build()
                
        except Exception as e:
            self._log(f'❌ 批处理失败: {e}')
            self._build_timer.stop()
            self._finish_build()
    
    def _finish_build(self):
        """完成构建"""
        try:
            # 恢复控件
            self.tree.setVisible(True)
            self.tree.blockSignals(False)
            self.tree.setUpdatesEnabled(True)
            
            # 立即隐藏加载状态，不显示任何完成信息
            self.refresh_status.hide_loading()
            
            # 清理
            if hasattr(self, '_build_items'):
                del self._build_items
            if hasattr(self, '_build_index'):
                del self._build_index
            if hasattr(self, '_notebooks_data'):
                del self._notebooks_data
                
        except Exception as e:
            self._log(f'❌ 完成构建时出错: {e}')

    def _find_item_by_id(self, id_: str):
        it = QTreeWidgetItemIterator(self.tree)
        while it.value():
            item = it.value()
            d=item.data(0,Qt.UserRole)
            if d and d.get('id')==id_: return item
            it+=1
        return None

    def _on_pop_done(self, nb:int, sec:int, pg:int):
        """完成界面构建"""
        # 立即隐藏加载状态
        self.refresh_status.hide_loading()
        self._log(f'✅ 读取完成：{nb} 笔记本，{sec} 分区，{pg} 页面')
        
        # 清理缓存
        if hasattr(self, '_item_cache'):
            self._item_cache.clear()
        self._set_busy(False)

    def _on_pop_err(self, msg:str):
        """构建失败处理"""
        self.refresh_status.hide_loading()
        self._set_busy(False)
        self._log(f'❌ 构建失败: {msg}')
        # 清理缓存
        if hasattr(self, '_item_cache'):
            self._item_cache.clear()

    def _on_item_changed(self, item, col):
        """处理树控件项目变化，实现级联勾选"""
        if col != 0:  # 只处理第一列的勾选变化
            return
            
        # 临时阻塞信号，避免级联操作触发无限递归
        self.tree.blockSignals(True)
        
        try:
            data = item.data(0, Qt.UserRole)
            if not data:
                return
                
            item_type = data.get('type')
            check_state = item.checkState(0)
            
            if item_type == 'notebook':
                # 勾选/取消勾选笔记本时，级联到所有分区和页面
                self._cascade_check_notebook(item, check_state)
            elif item_type == 'section':
                # 勾选/取消勾选分区时，级联到该分区下的所有页面
                self._cascade_check_section(item, check_state)
            elif item_type == 'page':
                # 页面勾选变化时，检查是否需要更新父分区的状态
                self._update_parent_check_state(item)
                
        finally:
            # 恢复信号
            self.tree.blockSignals(False)
            # 更新选择状态和转换按钮
            self._update_selection()
            self._update_convert()
    
    def _cascade_check_notebook(self, notebook_item, check_state):
        """级联勾选笔记本下的所有分区和页面"""
        for i in range(notebook_item.childCount()):
            section_item = notebook_item.child(i)
            section_item.setCheckState(0, check_state)
            # 级联到该分区下的所有页面
            self._cascade_check_section(section_item, check_state)
    
    def _cascade_check_section(self, section_item, check_state):
        """级联勾选分区下的所有页面"""
        for i in range(section_item.childCount()):
            page_item = section_item.child(i)
            page_item.setCheckState(0, check_state)
    
    def _update_parent_check_state(self, page_item):
        """根据子页面的勾选状态更新父分区的勾选状态"""
        section_item = page_item.parent()
        if not section_item:
            return
            
        # 检查分区下所有页面的勾选状态
        checked_count = 0
        total_count = section_item.childCount()
        
        for i in range(total_count):
            child = section_item.child(i)
            if child.checkState(0) == Qt.Checked:
                checked_count += 1
        
        # 根据子页面状态设置分区状态
        if checked_count == 0:
            section_item.setCheckState(0, Qt.Unchecked)
        elif checked_count == total_count:
            section_item.setCheckState(0, Qt.Checked)
        else:
            section_item.setCheckState(0, Qt.PartiallyChecked)
        
        # 递归更新笔记本状态
        self._update_notebook_check_state(section_item)
    
    def _update_notebook_check_state(self, section_item):
        """根据分区状态更新笔记本的勾选状态"""
        notebook_item = section_item.parent()
        if not notebook_item:
            return
            
        # 检查笔记本下所有分区的勾选状态
        checked_count = 0
        partial_count = 0
        total_count = notebook_item.childCount()
        
        for i in range(total_count):
            child = notebook_item.child(i)
            child_state = child.checkState(0)
            if child_state == Qt.Checked:
                checked_count += 1
            elif child_state == Qt.PartiallyChecked:
                partial_count += 1
        
        # 根据分区状态设置笔记本状态
        if checked_count == 0 and partial_count == 0:
            notebook_item.setCheckState(0, Qt.Unchecked)
        elif checked_count == total_count:
            notebook_item.setCheckState(0, Qt.Checked)
        else:
            notebook_item.setCheckState(0, Qt.PartiallyChecked)

    def _update_selection(self):
        sel=[]; it=QTreeWidgetItemIterator(self.tree)
        while it.value():
            item=it.value(); d=item.data(0,Qt.UserRole)
            if d and d.get('type')=='page' and item.checkState(0)==Qt.Checked:
                # 收集父级名称
                sec=item.parent(); nb=sec.parent() if sec else None
                sel.append({'page_id': d['id'], 'page_name': d['name'], 'section_name': (sec.data(0,Qt.UserRole) or {}).get('name',''), 'notebook_name': (nb.data(0,Qt.UserRole) or {}).get('name','')})
            it+=1
        self.selected_items=sel

    def _update_convert(self):
        ok = bool(self.selected_items) and bool(self.output_dir)
        self.convert_btn.setEnabled(ok)
        self.convert_btn.setText(f'🚀 开始转换 ({len(self.selected_items)} 个页面)' if self.selected_items else '🚀 开始转换')

    def _choose_dir(self):
        d = QFileDialog.getExistingDirectory(self, '选择输出目录', self.output_dir or os.path.expanduser('~'))
        if d:
            self.output_dir=d; self.lbl_out.setText(d); self._update_convert()

    def _select_all(self):
        it=QTreeWidgetItemIterator(self.tree)
        while it.value():
            item=it.value()
            if item.flags() & Qt.ItemIsUserCheckable:
                item.setCheckState(0,Qt.Checked)
            it+=1
        self._update_selection(); self._update_convert()

    def _select_none(self):
        it=QTreeWidgetItemIterator(self.tree)
        while it.value():
            item=it.value()
            if item.flags() & Qt.ItemIsUserCheckable:
                item.setCheckState(0,Qt.Unchecked)
            it+=1
        self._update_selection(); self._update_convert()

    def _convert(self):
        if not self.selected_items or not self.output_dir:
            QMessageBox.warning(self,'提示','请选择页面和输出目录')
            return
        
        self._set_busy(True)
        self.conv_status.show_loading('🚀 正在转换...')
        self.progress.setVisible(True)
        self.progress.setValue(0)
        self.log.clear()
        
        # 确保之前的转换线程已停止
        if self._convert_thread and self._convert_thread.isRunning():
            self._convert_thread.terminate()
            self._convert_thread.wait(100)
        
        self._convert_thread = _ConvertWorker(
            self.onenote, self.parser, self.selected_items, self.output_dir,
            self.cb_pdf.isChecked(), self.cb_docx.isChecked(),
            self.cb_img.isChecked(), self.cb_att.isChecked()
        )
        
        self._convert_thread.progress.connect(self.progress.setValue, Qt.QueuedConnection)
        self._convert_thread.msg.connect(self._log, Qt.QueuedConnection)
        self._convert_thread.done.connect(self._conv_done, Qt.QueuedConnection)
        self._convert_thread.err.connect(self._conv_err, Qt.QueuedConnection)
        self._convert_thread.start(QThread.NormalPriority)

    def _conv_done(self):
        self.progress.setValue(100); self.conv_status.hide_loading(); self._set_busy(False)
        QMessageBox.information(self,'完成','转换完成')

    def _conv_err(self, m:str):
        self.progress.setVisible(False); self.conv_status.hide_loading(); self._set_busy(False)
        QMessageBox.critical(self,'错误', m)
    
    def closeEvent(self, event):
        """关闭事件处理"""
        try:
            # 停止所有线程
            for thread in [getattr(self, '_loading_thread', None), 
                          getattr(self, '_populate_thread', None), 
                          getattr(self, '_convert_thread', None)]:
                if thread and thread.isRunning():
                    thread.terminate()
                    thread.wait(100)
            
            # 停止所有定时器
            if hasattr(self, '_build_timer') and self._build_timer.isActive():
                self._build_timer.stop()
            
            # 清理资源
            if hasattr(self, 'parser'):
                self.parser.cleanup_temp_files()
                
        except Exception:
            pass  # 忽略关闭时的错误
        finally:
            event.accept()
    
    def resizeEvent(self, event):
        """窗口大小改变时重新调整列宽"""
        super().resizeEvent(event)
        # 延迟调整列宽，确保窗口大小调整完成
        QTimer.singleShot(10, self._setup_tree_columns)


# ======= 线程 =======
class _DetectWorker(QThread):
    progress = pyqtSignal(str)
    done = pyqtSignal(dict)
    err = pyqtSignal(str)
    
    def __init__(self, api: OneNoteAPI):
        super().__init__()
        self.api=api
        self.setTerminationEnabled(True)
        
    def run(self):
        try:
            # 发送初始进度
            self.progress.emit('🔍 正在连接OneNote...')
            self.msleep(10)  # 让UI有机会更新
            
            if not self.api.initialize():
                self.err.emit('无法连接OneNote')
                return
            
            # 获取笔记本
            self.progress.emit('📚 正在获取笔记本列表...')
            self.msleep(10)  # 让UI有机会更新
            
            nbs = self.api.get_notebooks()
            if not nbs:
                self.err.emit('未发现笔记本')
                return
            
            # 计算统计信息
            total = sum(len(s.get('pages',{})) for nb in nbs.values() for s in nb.get('sections',{}).values())
            self.progress.emit(f'✅ 发现 {len(nbs)} 个笔记本，{total} 个页面')
            self.msleep(10)  # 让UI有机会更新
            
            self.done.emit(nbs)
        except Exception as e:
            self.err.emit(str(e))


class _PopulateWorker(QThread):
    all_data = pyqtSignal(dict)  # 一次性发送所有数据
    progress = pyqtSignal(int)
    msg = pyqtSignal(str)
    done = pyqtSignal(int,int,int)
    err = pyqtSignal(str)
    
    def __init__(self, notebooks: dict):
        super().__init__()
        self.nbs = notebooks
        self.setTerminationEnabled(True)
        
    def run(self):
        """一次性处理所有数据，不分批"""
        try:
            nb_count = len(self.nbs)
            sec_count = 0
            pg_count = 0
            
            # 统计数量
            for nb_data in self.nbs.values():
                for sec_data in nb_data.get('sections', {}).values():
                    sec_count += 1
                    pg_count += len(sec_data.get('pages', {}))
            
            # 发送进度更新
            self.progress.emit(50)
            self.msleep(10)  # 让UI有机会更新
            
            # 一次性发送所有数据，让UI线程处理
            self.all_data.emit(self.nbs)
            
            self.progress.emit(100)
            self.done.emit(nb_count, sec_count, pg_count)
            
        except Exception as e:
            self.err.emit(str(e))


class _ConvertWorker(QThread):
    progress = pyqtSignal(int)
    msg = pyqtSignal(str)
    done = pyqtSignal()
    err = pyqtSignal(str)
    def __init__(self, api: OneNoteAPI, parser: OneNoteContentParser, items: List[dict], out_dir: str, pdf: bool, docx: bool, images: bool, attachments: bool):
        super().__init__(); self.api=api; self.parser=parser; self.items=items; self.out=Path(out_dir); self.pdf=pdf; self.docx=docx; self.images=images; self.attach=attachments
    def run(self):
        try:
            n=len(self.items); done=0
            for it in self.items:
                pid=it['page_id']; name=it['page_name']; nb=it['notebook_name']; sec=it['section_name']
                safe=lambda s: ''.join(c for c in (s or '未命名') if c.isalnum() or c in (' ','-','_','.')).strip()[:100] or '未命名'
                d = self.out/safe(nb)/safe(sec); d.mkdir(parents=True, exist_ok=True)
                xml = self.api.get_page_content(pid)
                if not xml: self.msg.emit(f'⚠️ 空页面: {name}'); continue
                
                # Word: 内嵌附件
                if self.docx:
                    out = d/f'{safe(name)}.docx'
                    att = d/f'{safe(name)}_attachments' if self.attach else None
                    ok = self.parser.parse_page_to_docx(xml, name, str(out), 
                                                        include_images=self.images, 
                                                        include_attachments=self.attach,
                                                        embed_attachments=True,  # Word内嵌附件
                                                        attachments_output_dir=att)
                    self.msg.emit(f'{"✅" if ok else "❌"} Word: {name}')
                
                # PDF: 附件保存到目录
                if self.pdf:
                    out = d/f'{safe(name)}.pdf'
                    att = d/f'{safe(name)}_attachments' if self.attach else None
                    ok = self.parser.parse_page_to_pdf(xml, name, str(out), 
                                                       include_images=self.images,
                                                       include_attachments=self.attach,
                                                       attachments_output_dir=att)  # PDF附件外置
                    self.msg.emit(f'{"✅" if ok else "❌"} PDF: {name}')
                
                done+=1; self.progress.emit(int(done/max(n,1)*100))
            self.done.emit()
        except Exception as e:
            self.err.emit(str(e))


def main():
    # Windows隐藏控制台：建议用 pythonw.exe 运行
    app = QApplication(sys.argv)
    
    # 设置应用程序名称
    app.setApplicationName('WATS-OneNote_Exporter')
    app.setOrganizationName('WATS')
    
    # 性能优化设置
    app.setAttribute(Qt.AA_EnableHighDpiScaling, True)  # 高DPI支持
    app.setAttribute(Qt.AA_UseHighDpiPixmaps, True)  # 高DPI图标
    app.setAttribute(Qt.AA_CompressHighFrequencyEvents, True)  # 压缩高频事件
    
    # 创建并显示主窗口
    w = ModernOneNoteGUI()
    w.show()
    
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()

