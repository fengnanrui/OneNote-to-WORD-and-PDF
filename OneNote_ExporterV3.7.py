#!/usr/bin/env pythonw
# -*- coding: utf-8 -*-

import sys
import os
import logging
import base64
import traceback
import ctypes
import tempfile
import subprocess
import re
import html
import time
import uuid
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
import xml.etree.ElementTree as ET

# ======= 依赖（尽量最少） =======
try:
    import win32com.client  # type: ignore
    import comtypes.client  # type: ignore
    COM_AVAILABLE = True
except Exception:
    COM_AVAILABLE = False

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QVBoxLayout, QHBoxLayout, QWidget,
    QPushButton, QLabel, QTreeWidget, QTreeWidgetItem, QProgressBar,
    QTextEdit, QGroupBox, QCheckBox, QFileDialog, QMessageBox, QSplitter,
    QTreeWidgetItemIterator
)
from PyQt5.QtCore import Qt, QTimer, QDateTime, QThread, pyqtSignal, QCoreApplication

# Word/PDF 依赖
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.section import WD_ORIENT, WD_SECTION

from reportlab.lib.pagesizes import A4
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
    Table, TableStyle, KeepInFrame, PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT


# ======= 工具函数 =======
def is_admin() -> bool:
    try:
        return ctypes.windll.shell32.IsUserAnAdmin()
    except Exception:
        return False


def check_onenote_process() -> bool:
    try:
        r = subprocess.run(['tasklist', '/FI', 'IMAGENAME eq ONENOTE.EXE'], capture_output=True, text=True)
        return 'ONENOTE.EXE' in r.stdout
    except Exception:
        return False


# ======= 一些轻量 UI 组件 =======
from PyQt5.QtCore import QTimer
from PyQt5.QtGui import QPainter, QPen, QColor


class LoadingIndicator(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedSize(26, 26)
        self.angle = 0
        self._timer = QTimer(self)
        self._timer.timeout.connect(self._tick)
        self._timer.setInterval(60)
        self.hide()

    def start(self, show_text=True):
        self._timer.start()
        self.show()

    def stop(self):
        self._timer.stop()
        self.hide()

    def _tick(self):
        self.angle = (self.angle + 10) % 360
        self.update()

    def paintEvent(self, e):
        p = QPainter(self)
        p.setRenderHint(QPainter.Antialiasing)
        pen = QPen(QColor(70, 130, 180))
        pen.setWidth(3)
        p.setPen(pen)
        r = 10
        p.translate(self.rect().center())
        p.rotate(self.angle)
        p.drawArc(-r, -r, 2*r, 2*r, 0, 120*16)


class StatusIndicator(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self._label = QLabel('', self)
        self._spinner = LoadingIndicator(self)
        layout = QHBoxLayout(self)
        layout.setContentsMargins(6, 6, 6, 6)
        layout.addWidget(self._spinner)
        layout.addWidget(self._label)
        self.hide()

    def show_loading(self, text: str):
        if self._label.text() != text:
            self._label.setText(text)
        if not self._spinner._timer.isActive():
            self._spinner.start()
        if not self.isVisible():
            self.show()

    def hide_loading(self):
        self._spinner.stop()
        self.hide()


# ======= OneNote API（COM优先，PowerShell回退） =======
class OneNoteAPI:
    def __init__(self):
        self.app = None
        self.temp_dir = tempfile.mkdtemp(prefix='onenote_app_')
        self.logger = logging.getLogger('OneNoteAPI')

    def initialize(self) -> bool:
        try:
            if not COM_AVAILABLE:
                raise RuntimeError('COM not available')
            admin = is_admin(); running = check_onenote_process()
            self.logger.info(f'权限: admin={admin}, running={running}')
            # 尝试三种COM
            try:
                self.app = win32com.client.gencache.EnsureDispatch('OneNote.Application')
                _ = self.app.GetHierarchy('', 1)
                return True
            except Exception as e:
                self.logger.warning(f'gencache失败: {e}')
            try:
                self.app = win32com.client.Dispatch('OneNote.Application')
                _ = self.app.GetHierarchy('', 1)
                return True
            except Exception as e:
                self.logger.warning(f'Dispatch失败: {e}')
            try:
                self.app = comtypes.client.CreateObject('OneNote.Application')
                _ = self.app.GetHierarchy('', 1)
                return True
            except Exception as e:
                self.logger.warning(f'comtypes失败: {e}')
            # 退到仅PS
            self.app = None
            return True
        except Exception as e:
            self.logger.error(f'初始化失败: {e}')
            return False

    def _ps(self, script: str) -> str:
        f = Path(self.temp_dir) / 'tmp.ps1'
        f.write_text(script, encoding='utf-8')
        r = subprocess.run(['powershell', '-ExecutionPolicy', 'Bypass', '-File', str(f)], capture_output=True, text=True, encoding='utf-8', errors='replace')
        return (r.stdout or '')

    def _get_hierarchy_ps(self, obj_id: str, scope: int) -> str:
        obj = obj_id.replace('"','""') if obj_id else ''
        sc = scope
        script = f"""
        [Console]::OutputEncoding=[System.Text.Encoding]::UTF8
        try {{
          $o=New-Object -ComObject OneNote.Application
          $x=""
          $o.GetHierarchy("{obj}",{sc},[ref]$x)
          Write-Output "SUCCESS:$x"
        }} catch {{ Write-Output "ERROR:$($_.Exception.Message)" }}
        """
        out = self._ps(script).strip()
        if out.startswith('SUCCESS:'): return out[8:]
        return ''

    def _get_page_ps(self, page_id: str) -> str:
        pid = page_id.replace('"','""')
        script = f"""
        [Console]::OutputEncoding=[System.Text.Encoding]::UTF8
        try {{
          $o=New-Object -ComObject OneNote.Application
          $x=""
          $o.GetPageContent("{pid}",[ref]$x,7)
          Write-Output "SUCCESS:$x"
        }} catch {{ Write-Output "ERROR:$($_.Exception.Message)" }}
        """
        out = self._ps(script).strip()
        if out.startswith('SUCCESS:'): return out[8:]
        return ''

    def get_notebooks(self) -> Dict:
        """获取笔记本列表，优化版本"""
        xml = ''
        try:
            if self.app:
                try:
                    xml = self.app.GetHierarchy('', 4)
                except Exception:
                    pass
            if not xml:
                xml = self._get_hierarchy_ps('', 4)
        except Exception:
            xml=''
        
        if not xml:
            return {}
        
        try:
            root = ET.fromstring(xml)
        except Exception:
            return {}
        
        def findall_local(p, name):
            return [e for e in p.iter() if (isinstance(e.tag,str) and (e.tag.endswith('}'+name) or e.tag==name or e.tag.split('}')[-1]==name))]
        
        notebooks={}
        for nb in findall_local(root, 'Notebook'):
            nb_id = nb.get('ID')
            nb_name = nb.get('name')
            if not nb_id or not nb_name: 
                continue
            
            sections={}
            for sec in findall_local(nb,'Section'):
                sid = sec.get('ID')
                sname = sec.get('name')
                if not sid or not sname: 
                    continue
                
                pages={}
                for pg in findall_local(sec,'Page'):
                    pid = pg.get('ID')
                    pname = pg.get('name')
                    if pid and pname:
                        pages[pid] = {'id':pid, 'name':pname}
                
                sections[sid] = {'id':sid, 'name':sname, 'pages':pages}
            
            notebooks[nb_id] = {'id':nb_id, 'name':nb_name, 'sections':sections}
        
        return notebooks

    def get_page_content(self, page_id: str, max_retries: int = 3) -> str:
        """获取页面内容，增加重试机制"""
        for attempt in range(max_retries):
            try:
                if self.app:
                    try:
                        c = self.app.GetPageContent(page_id, 7)
                        if c and c.strip(): 
                            return c
                    except Exception:
                        pass
                    try:
                        x=''; self.app.GetPageContent(page_id, x, 7)
                        if x and x.strip(): 
                            return x
                    except Exception:
                        pass
                content = self._get_page_ps(page_id)
                if content and content.strip():
                    return content
            except Exception as e:
                self.logger.warning(f"获取页面内容失败 (尝试 {attempt+1}/{max_retries}): {e}")
                if attempt == max_retries - 1:
                    self.logger.error(f"无法获取页面 {page_id} 的内容: {e}")
                else:
                    time.sleep(1)
        return ''


# ======= 增强的解析器（Word / PDF） =======
class EnhancedOneNoteContentParser:
    def __init__(self):
        self.logger = logging.getLogger('EnhancedParser')
        self.temp_files: List[str] = []
        self._setup_chinese_fonts()
        
        # 增强的图片格式支持
        self.supported_image_extensions = {
            b'\xFF\xD8\xFF': '.jpg',
            b'\x89PNG\r\n\x1a\n': '.png', 
            b'GIF87a': '.gif',
            b'GIF89a': '.gif',
            b'\x00\x00\x01\x00': '.ico',
            b'BM': '.bmp',
            b'RIFF': '.webp'
        }
    
    def _setup_chinese_fonts(self):
        """设置中文字体支持"""
        try:
            chinese_fonts = [
                ('SimSun', 'C:/Windows/Fonts/simsun.ttc'),
                ('SimHei', 'C:/Windows/Fonts/simhei.ttf'), 
                ('Microsoft YaHei', 'C:/Windows/Fonts/msyh.ttc'),
                ('PingFang SC', 'C:/Windows/Fonts/PingFang.ttc'),
            ]
            
            self.chinese_font = None
            for font_name, font_path in chinese_fonts:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        self.chinese_font = font_name
                        self.logger.info(f"成功注册中文字体: {font_name}")
                        break
                    except Exception as e:
                        self.logger.debug(f"注册字体{font_name}失败: {e}")
                        continue
            
            if not self.chinese_font:
                self.chinese_font = 'Helvetica'
                self.logger.warning("未找到中文字体，使用默认字体")
                
        except Exception as e:
            self.chinese_font = 'Helvetica'
            self.logger.error(f"字体设置失败: {e}")
    
    def cleanup_temp_files(self):
        """清理临时文件"""
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception as e:
                self.logger.debug(f"清理临时文件失败: {e}")
        self.temp_files.clear()

    def _findall_local(self, parent: ET.Element, local_name: str) -> List[ET.Element]:
        """命名空间无关的元素查找"""
        out=[]
        for el in parent.iter():
            tag = el.tag
            if isinstance(tag,str) and (tag.endswith('}'+local_name) or tag==local_name or tag.split('}')[-1]==local_name):
                if el is not parent or tag.split('}')[-1] != local_name:
                    out.append(el)
        return out

    def _is_inside_element(self, element: ET.Element, target_elements: List[str]) -> bool:
        """检查元素是否位于指定元素内部"""
        current = element
        while current is not None:
            tag_name = current.tag
            if isinstance(tag_name, str):
                for target in target_elements:
                    if (tag_name.endswith('}'+target) or tag_name == target):
                        return True
            current = current.getparent() if hasattr(current, 'getparent') else None
        return False

    def _extract_image_data_enhanced(self, img_elem: ET.Element) -> Optional[Tuple[bytes, str]]:
        """增强的图片数据提取，支持多种格式和属性"""
        
        # 扩展的属性列表
        image_attributes = [
            'data', 'Data', 'binaryData', 'base64Data', 'imageData',
            'src', 'source', 'content', 'bytes', 'binary'
        ]
        
        # 1. 从元素属性中提取
        for attr in image_attributes:
            value = img_elem.get(attr)
            if value:
                try:
                    data = base64.b64decode(value)
                    if len(data) > 100:  # 有效的图片数据应该大于100字节
                        format_ext = self._detect_image_format(data)
                        self.logger.debug(f"从属性 {attr} 提取到图片数据: {len(data)} bytes, 格式: {format_ext}")
                        return data, format_ext
                except Exception as e:
                    self.logger.debug(f"解码属性 {attr} 失败: {e}")
                    continue
        
        # 2. 从子元素中提取
        for child in img_elem:
            if child.text:
                # 检查子元素标签
                tag = child.tag
                if isinstance(tag, str) and any(keyword in tag.lower() for keyword in ['data', 'binary', 'content']):
                    try:
                        data = base64.b64decode(child.text)
                        if len(data) > 100:
                            format_ext = self._detect_image_format(data)
                            self.logger.debug(f"从子元素 {tag} 提取到图片数据: {len(data)} bytes, 格式: {format_ext}")
                            return data, format_ext
                    except Exception as e:
                        self.logger.debug(f"解码子元素 {tag} 失败: {e}")
                        continue
        
        # 3. 递归搜索所有后代元素
        for descendant in img_elem.iter():
            if descendant != img_elem and descendant.text:
                # 尝试解析任何可能包含base64数据的文本
                text = descendant.text.strip()
                if len(text) > 100 and self._looks_like_base64(text):
                    try:
                        data = base64.b64decode(text)
                        if len(data) > 100:
                            format_ext = self._detect_image_format(data)
                            self.logger.debug(f"从后代元素递归提取到图片数据: {len(data)} bytes, 格式: {format_ext}")
                            return data, format_ext
                    except Exception:
                        continue
        
        self.logger.debug("未找到有效的图片数据")
        return None
    
    def _looks_like_base64(self, text: str) -> bool:
        """检查文本是否看起来像base64编码"""
        if len(text) < 100:  # 太短的不太可能是图片
            return False
        # base64字符集检查
        import string
        valid_chars = string.ascii_letters + string.digits + '+/='
        return all(c in valid_chars for c in text[:100])  # 只检查前100个字符以提高效率
    
    def _detect_image_format(self, data: bytes) -> str:
        """检测图片格式"""
        for signature, ext in self.supported_image_extensions.items():
            if data.startswith(signature):
                return ext
        
        # 额外检查WEBP（RIFF格式的特殊情况）
        if data.startswith(b'RIFF') and b'WEBP' in data[:12]:
            return '.webp'
        
        return '.png'  # 默认使用PNG格式

    def _extract_table_data_enhanced(self, table_elem: ET.Element) -> List[List[str]]:
        """增强的表格数据提取，处理复杂结构和避免重复"""
        rows = []
        seen_row_signatures = set()  # 用于去重
        
        # 查找所有行元素
        row_elements = self._findall_local(table_elem, 'Row')
        
        for row_elem in row_elements:
            cell_elements = self._findall_local(row_elem, 'Cell')
            row_data = []
            
            for cell_elem in cell_elements:
                # 使用增强的单元格文本提取
                cell_text = self._extract_cell_text_enhanced(cell_elem)
                
                # 处理合并单元格属性
                colspan = self._get_cell_span(cell_elem, 'colspan')
                rowspan = self._get_cell_span(cell_elem, 'rowspan') 
                
                # 添加单元格数据
                row_data.append(cell_text)
                
                # 如果有列合并，添加空单元格占位
                for _ in range(colspan - 1):
                    row_data.append('')
            
            # 创建行的签名用于去重（基于前3个单元格的内容）
            row_signature = '|'.join(row_data[:3]) if len(row_data) >= 3 else '|'.join(row_data)
            
            # 只添加非空且未见过的行
            if row_data and any(cell.strip() for cell in row_data) and row_signature not in seen_row_signatures:
                seen_row_signatures.add(row_signature)
                rows.append(row_data)
                self.logger.debug(f"提取表格行: {len(row_data)} 个单元格")
        
        self.logger.info(f"表格数据提取完成: {len(rows)} 行, 最大列数: {max(len(row) for row in rows) if rows else 0}")
        return rows
    
    def _get_cell_span(self, cell_elem: ET.Element, span_type: str) -> int:
        """获取单元格的跨度属性"""
        try:
            span_attrs = [span_type, span_type.upper(), span_type.capitalize()]
            for attr in span_attrs:
                value = cell_elem.get(attr)
                if value:
                    return int(value)
        except (ValueError, TypeError):
            pass
        return 1
    
    def _extract_cell_text_enhanced(self, cell_elem: ET.Element) -> str:
        """增强的单元格文本提取，处理嵌套内容"""
        text_parts = []
        seen_texts = set()
        
        # 1. 查找所有文本元素（T元素）
        for t_elem in self._findall_local(cell_elem, 'T'):
            if t_elem.text:
                clean_text = self._clean_text_content(t_elem.text)
                if clean_text and clean_text not in seen_texts:
                    seen_texts.add(clean_text)
                    text_parts.append(clean_text)
        
        # 2. 如果没有找到T元素，递归提取所有文本内容
        if not text_parts:
            def collect_text_recursive(elem):
                if elem.text and elem.text.strip():
                    clean_text = self._clean_text_content(elem.text)
                    if clean_text and clean_text not in seen_texts:
                        seen_texts.add(clean_text)
                        text_parts.append(clean_text)
                
                for child in elem:
                    collect_text_recursive(child)
                    if child.tail and child.tail.strip():
                        clean_text = self._clean_text_content(child.tail)
                        if clean_text and clean_text not in seen_texts:
                            seen_texts.add(clean_text)
                            text_parts.append(clean_text)
            
            collect_text_recursive(cell_elem)
        
        # 3. 合并文本，保持适当的间距
        result = ' '.join(text_parts)
        result = re.sub(r'\s+', ' ', result).strip()
        
        return result
    
    def _clean_text_content(self, text: str) -> str:
        """清理文本内容"""
        if not text:
            return ""
        
        # HTML解码
        text = html.unescape(text)
        # 移除HTML标签
        text = re.sub(r'<[^>]+>', '', text)
        # 处理特殊字符
        text = text.replace('\u2022', '•').replace('\u2013', '-').replace('\u2014', '—')
        text = text.replace('\u201c', '"').replace('\u201d', '"')
        text = text.replace('\u2018', "'").replace('\u2019', "'")
        # 清理空白字符
        text = re.sub(r'[\t\x0b\x0c]+', ' ', text)
        text = text.strip()
        
        return text

    def _process_content_in_original_order(self, root: ET.Element, processor_func, *args) -> None:
        """按照XML中的原始顺序处理所有内容元素"""
        
        def process_element(elem: ET.Element, depth: int = 0):
            """深度优先遍历处理元素"""
            tag_name = elem.tag
            if not isinstance(tag_name, str):
                return
            
            # 获取标签的本地名称
            local_name = tag_name.split('}')[-1] if '}' in tag_name else tag_name
            
            # 根据元素类型调用对应的处理器
            if local_name == 'Table':
                processor_func('table', elem, *args)
            elif local_name == 'Image':  
                processor_func('image', elem, *args)
            elif local_name in ['OE', 'T'] and not self._is_inside_element(elem, ['Table']):
                # 文本元素，但不在表格内
                processor_func('text', elem, *args)
            
            # 递归处理子元素
            for child in elem:
                process_element(child, depth + 1)
        
        # 从根元素开始处理
        for child in root:
            process_element(child)

    # === Word处理方法 ===
    def parse_page_to_docx(self, xml: str, page_name: str, out_path: str,
                           include_images=True, include_attachments=True,
                           embed_attachments=False,
                           attachments_output_dir: Optional[Path]=None) -> bool:
        try:
            root = ET.fromstring(xml)
            doc = Document()
            doc.add_heading(page_name, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 按原始顺序处理内容
            self._process_content_in_original_order(root, self._word_content_processor, 
                                                    doc, include_images)

            doc.save(out_path)
            return True
        except Exception as e:
            self.logger.error(f'DOCX生成失败: {e}')
            traceback.print_exc()
            return False
    
    def _word_content_processor(self, element_type: str, element: ET.Element, 
                               doc: Document, include_images: bool):
        """Word内容处理器"""
        try:
            if element_type == 'table':
                self._process_table_for_word(element, doc)
            elif element_type == 'image' and include_images:
                self._process_image_for_word(element, doc)  
            elif element_type == 'text':
                self._process_text_for_word(element, doc)
        except Exception as e:
            self.logger.warning(f"处理{element_type}元素失败: {e}")
    
    def _process_table_for_word(self, table_elem: ET.Element, doc: Document):
        """处理Word表格"""
        rows_data = self._extract_table_data_enhanced(table_elem)
        if not rows_data:
            return
            
        max_cols = max(len(row) for row in rows_data) if rows_data else 1
        max_rows = len(rows_data)
        
        try:
            # 创建表格
            wt = doc.add_table(rows=max_rows, cols=max_cols)
            wt.style = 'Table Grid'
            wt.autofit = True
            
            # 填充数据
            for i, row_data in enumerate(rows_data):
                for j, cell_text in enumerate(row_data):
                    if i < len(wt.rows) and j < len(wt.rows[i].cells):
                        cell = wt.rows[i].cells[j]
                        cell.text = cell_text or ''
                        
                        # 设置单元格格式
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.word_wrap = True
                            paragraph.paragraph_format.keep_together = True
            
            doc.add_paragraph()  # 表格后添加空行
            
        except Exception as e:
            self.logger.warning(f"创建Word表格失败: {e}")
    
    def _process_image_for_word(self, img_elem: ET.Element, doc: Document):
        """处理Word图片"""
        image_data = self._extract_image_data_enhanced(img_elem)
        if not image_data:
            return
            
        data, format_ext = image_data
        
        # 创建临时图片文件
        temp_img = tempfile.mktemp(suffix=format_ext)
        self.temp_files.append(temp_img)
        
        try:
            Path(temp_img).write_bytes(data)
            
            # 智能调整图片尺寸
            display_width = self._calculate_word_image_width(data)
            doc.add_picture(temp_img, width=Inches(display_width))
            doc.add_paragraph()
            
            self.logger.debug(f"添加Word图片成功: {len(data)} bytes, 宽度: {display_width}英寸")
            
        except Exception as e:
            self.logger.warning(f"添加Word图片失败: {e}")
    
    def _calculate_word_image_width(self, image_data: bytes) -> float:
        """计算Word文档中的图片显示宽度"""
        try:
            # 尝试获取图片尺寸
            try:
                from PIL import Image as PILImage
                import io
                with PILImage.open(io.BytesIO(image_data)) as pil_img:
                    orig_width, orig_height = pil_img.size
                    aspect_ratio = orig_height / orig_width
            except ImportError:
                # 没有PIL时使用默认比例
                aspect_ratio = 0.75
                orig_width = 800
            
            # Word页面可用宽度（约6.5英寸）
            max_width = 6.5
            min_width = 2.0
            
            # 根据原始宽度智能选择显示宽度
            if orig_width <= 400:
                # 小图片，适度放大
                display_width = min(max_width * 0.6, 4.0)
            elif orig_width <= 800:
                # 中等图片
                display_width = min(max_width * 0.8, 5.0)
            elif orig_width <= 1600:
                # 大图片
                display_width = max_width * 0.9
            else:
                # 超大图片，使用最大宽度
                display_width = max_width
            
            # 如果图片很高，限制宽度以防止过高
            if aspect_ratio > 1.5:  # 高图片
                display_width = min(display_width, max_width * 0.7)
            
            return max(min_width, display_width)
            
        except Exception:
            return 4.0  # 默认宽度
    
    def _process_text_for_word(self, text_elem: ET.Element, doc: Document):
        """处理Word文本"""
        if text_elem.text:
            text = self._clean_text_content(text_elem.text)
            if text.strip():
                p = doc.add_paragraph()
                
                # 获取缩进级别
                indent_level = self._get_text_indent_level(text_elem)
                if indent_level > 0:
                    p.paragraph_format.left_indent = Inches(indent_level * 0.25)
                
                # 添加文本运行并应用格式
                run = p.add_run(text)
                self._apply_text_formatting_word(text_elem, run)
    
    def _get_text_indent_level(self, text_elem: ET.Element) -> int:
        """获取文本缩进级别"""
        # 查找父级List元素
        current = text_elem
        while current is not None:
            if hasattr(current, 'getparent'):
                parent = current.getparent()
                if parent is not None:
                    tag_name = parent.tag
                    if isinstance(tag_name, str) and (tag_name.endswith('}List') or tag_name == 'List'):
                        try:
                            return int(parent.get('indent', '0'))
                        except (ValueError, TypeError):
                            return 0
                current = parent
            else:
                break
        return 0
    
    def _apply_text_formatting_word(self, elem: ET.Element, run):
        """应用Word文本格式"""
        try:
            # 检查格式属性
            if elem.get('bold') == 'true' or 'bold' in elem.tag.lower():
                run.bold = True
            if elem.get('italic') == 'true' or 'italic' in elem.tag.lower():
                run.italic = True
            if elem.get('underline') == 'true' or 'underline' in elem.tag.lower():
                run.underline = True
            
            # 字体大小
            font_size = elem.get('fontSize')
            if font_size:
                try:
                    run.font.size = Pt(float(font_size))
                except (ValueError, TypeError):
                    pass
                    
        except Exception as e:
            self.logger.debug(f"应用Word格式失败: {e}")

    # === PDF处理方法 ===
    def parse_page_to_pdf(self, xml: str, page_name: str, out_path: str,
                          include_images=True, include_attachments=True,
                          attachments_output_dir: Optional[Path]=None) -> bool:
        try:
            root = ET.fromstring(xml)
            
            # 创建自定义样式，支持中文
            styles = getSampleStyleSheet()
            
            title_style = ParagraphStyle(
                'ChineseTitle',
                parent=styles['Heading1'],
                fontSize=18,
                alignment=TA_CENTER,
                fontName=self.chinese_font,
                textColor=colors.black,
                spaceAfter=12
            )
            
            normal_style = ParagraphStyle(
                'ChineseNormal',
                parent=styles['Normal'],
                fontSize=12,
                fontName=self.chinese_font,
                textColor=colors.black,
                leftIndent=0,
                rightIndent=0,
                spaceAfter=6
            )
            
            # 创建文档
            doc = SimpleDocTemplate(
                out_path, 
                pagesize=A4,
                leftMargin=1.5*cm,
                rightMargin=1.5*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )
            
            story = []
            story.append(Paragraph(page_name, title_style))
            story.append(Spacer(1, 12))

            # 按原始顺序处理内容
            self._process_content_in_original_order(root, self._pdf_content_processor, 
                                                    story, normal_style, include_images)

            doc.build(story)
            return True
        except Exception as e:
            self.logger.error(f'PDF生成失败: {e}')
            traceback.print_exc()
            return False
    
    def _pdf_content_processor(self, element_type: str, element: ET.Element,
                              story: List, normal_style: ParagraphStyle, include_images: bool):
        """PDF内容处理器"""
        try:
            if element_type == 'table':
                self._process_table_for_pdf(element, story, normal_style)
            elif element_type == 'image' and include_images:
                self._process_image_for_pdf(element, story)
            elif element_type == 'text':
                self._process_text_for_pdf(element, story, normal_style)
        except Exception as e:
            self.logger.warning(f"处理PDF {element_type}元素失败: {e}")
    
    def _process_table_for_pdf(self, table_elem: ET.Element, story: List, normal_style: ParagraphStyle):
        """处理PDF表格"""
        rows_data = self._extract_table_data_enhanced(table_elem)
        if not rows_data:
            return
            
        # 创建表格样式
        cell_style = ParagraphStyle(
            'TableCell',
            parent=normal_style,
            fontSize=9,
            leading=11,
            fontName=normal_style.fontName,
            leftIndent=2,
            rightIndent=2,
            spaceAfter=2,
            spaceBefore=2
        )
        
        max_cols = max(len(row) for row in rows_data) if rows_data else 1
        
        # 补齐行数据
        normalized_rows = []
        for row in rows_data:
            normalized_row = row + [''] * (max_cols - len(row))
            normalized_rows.append(normalized_row)
        
        try:
            # 转换为Paragraph对象
            table_flow = []
            for row in normalized_rows:
                flow_row = []
                for cell_text in row:
                    # 限制单元格文本长度
                    if len(cell_text) > 150:
                        cell_text = cell_text[:147] + '...'
                    para = Paragraph(cell_text or ' ', cell_style)
                    flow_row.append(para)
                table_flow.append(flow_row)
            
            if table_flow:
                # 计算合适的列宽
                available_width = A4[0] - 3*cm
                col_width = available_width / max_cols
                col_width = max(col_width, 1.2*cm)  # 最小列宽
                col_widths = [col_width] * max_cols
                
                pdf_table = Table(table_flow, colWidths=col_widths)
                
                table_style = TableStyle([
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('FONTNAME', (0, 0), (-1, -1), normal_style.fontName),
                    ('LEFTPADDING', (0, 0), (-1, -1), 2),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 2),
                    ('TOPPADDING', (0, 0), (-1, -1), 2),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ])
                
                pdf_table.setStyle(table_style)
                
                story.append(Spacer(1, 6))
                story.append(pdf_table)
                story.append(Spacer(1, 12))
                
                self.logger.debug(f"添加PDF表格成功: {len(normalized_rows)}行 x {max_cols}列")
                
        except Exception as e:
            self.logger.warning(f"PDF表格渲染失败: {e}")
    
    def _process_image_for_pdf(self, img_elem: ET.Element, story: List):
        """处理PDF图片"""
        image_data = self._extract_image_data_enhanced(img_elem)
        if not image_data:
            return
            
        data, format_ext = image_data
        
        # 创建临时图片文件
        temp_img = tempfile.mktemp(suffix=format_ext)
        self.temp_files.append(temp_img)
        
        try:
            Path(temp_img).write_bytes(data)
            
            # 计算合适的显示尺寸
            width, height = self._calculate_pdf_image_size(data)
            
            img = RLImage(temp_img, width=width, height=height)
            story.append(Spacer(1, 8))
            story.append(img)
            story.append(Spacer(1, 12))
            
            self.logger.debug(f"添加PDF图片成功: {len(data)} bytes, 尺寸: {width}x{height}")
            
        except Exception as e:
            self.logger.warning(f"处理PDF图片失败: {e}")
    
    def _calculate_pdf_image_size(self, image_data: bytes) -> Tuple[float, float]:
        """计算PDF中的图片显示尺寸"""
        try:
            try:
                from PIL import Image as PILImage
                import io
                with PILImage.open(io.BytesIO(image_data)) as pil_img:
                    orig_width, orig_height = pil_img.size
            except ImportError:
                orig_width, orig_height = 600, 400
            
            # 计算合适的显示尺寸
            page_width = A4[0] - 3*cm
            page_height = A4[1] - 4*cm
            
            # 智能缩放
            scale_w = page_width / orig_width
            scale_h = page_height / orig_height
            scale = min(scale_w, scale_h, 1.0)  # 不放大
            
            # 确保最小尺寸
            final_width = max(orig_width * scale, page_width * 0.3)
            final_height = max(orig_height * scale, 100)
            
            # 确保不超过页面尺寸
            final_width = min(final_width, page_width)
            final_height = min(final_height, page_height * 0.8)
            
            return final_width, final_height
            
        except Exception:
            return 4*inch, 3*inch  # 默认尺寸
    
    def _process_text_for_pdf(self, text_elem: ET.Element, story: List, normal_style: ParagraphStyle):
        """处理PDF文本"""
        if text_elem.text:
            text = self._clean_text_content(text_elem.text)
            if text.strip():
                # 获取缩进级别
                indent_level = self._get_text_indent_level(text_elem)
                
                # 创建带缩进的样式
                if indent_level > 0:
                    text_style = ParagraphStyle(
                        f'Indent{indent_level}',
                        parent=normal_style,
                        leftIndent=indent_level * 15,
                        bulletIndent=indent_level * 10 if indent_level > 0 else 0
                    )
                else:
                    text_style = normal_style
                
                story.append(Paragraph(text, text_style))
                story.append(Spacer(1, 3))


# ======= GUI =======
class ModernOneNoteGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.onenote = OneNoteAPI()
        self.parser = EnhancedOneNoteContentParser()  # 使用增强解析器
        self.selected_items=[]; self.output_dir=''
        self._busy=False
        self._loading_thread = None
        self._populate_thread = None
        self._convert_thread = None
        self._setup_logging(); self._init_ui(); self._apply_styles()
        
        self.setAttribute(Qt.WA_OpaquePaintEvent, True)
        self.setAttribute(Qt.WA_NoSystemBackground, True)
        
        QTimer.singleShot(500, self._auto_detect)

    def _setup_logging(self):
        logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

    def _init_ui(self):
        self.setWindowTitle('WATS-OneNote_Exporter')
        self.setMinimumSize(1000, 620)
        
        cw = QWidget()
        self.setCentralWidget(cw)
        main = QHBoxLayout(cw)
        
        # 左侧面板
        left = QWidget()
        lv = QVBoxLayout(left)
        
        title = QLabel('📚 OneNote 笔记本导出工具')
        title.setObjectName('title_label')
        lv.addWidget(title)
        
        self.refresh_btn = QPushButton('🔄 刷新笔记本')
        self.refresh_btn.clicked.connect(self._refresh)
        lv.addWidget(self.refresh_btn)
        
        self.refresh_status = StatusIndicator(left)
        self.refresh_status.setObjectName('status_bar')
        lv.addWidget(self.refresh_status)
        
        self.tree = QTreeWidget()
        self.tree.setObjectName('notebook_tree')
        self.tree.setHeaderLabels(['笔记本/分区/页面','状态'])
        self.tree.itemChanged.connect(self._on_item_changed)
        
        header = self.tree.header()
        header.setStretchLastSection(False)
        QTimer.singleShot(100, self._setup_tree_columns)
        
        self.tree.setUniformRowHeights(True)
        self.tree.setAlternatingRowColors(False)
        self.tree.setAnimated(False)
        self.tree.setExpandsOnDoubleClick(True)
        self.tree.setItemsExpandable(True)
        self.tree.setRootIsDecorated(True)
        self.tree.setIndentation(20)
        
        lv.addWidget(self.tree)
        
        # 选择按钮
        sel_bar = QWidget()
        hb = QHBoxLayout(sel_bar)
        self.btn_all = QPushButton('✅ 全选')
        self.btn_all.clicked.connect(self._select_all)
        self.btn_none = QPushButton('❌ 取消全选')
        self.btn_none.clicked.connect(self._select_none)
        hb.addWidget(self.btn_all)
        hb.addWidget(self.btn_none)
        lv.addWidget(sel_bar)

        # 右侧面板
        right = QWidget()
        rv = QVBoxLayout(right)
        
        # 输出设置
        out_g = QGroupBox('📁 输出设置')
        out_g.setObjectName('group')
        og = QHBoxLayout(out_g)
        self.lbl_out = QLabel('未选择输出目录')
        self.btn_dir = QPushButton('选择目录')
        self.btn_dir.clicked.connect(self._choose_dir)
        og.addWidget(self.lbl_out)
        og.addWidget(self.btn_dir)
        rv.addWidget(out_g)
        
        # 导出格式
        fmt_g = QGroupBox('📄 导出格式')
        fmt_g.setObjectName('group')
        fg = QVBoxLayout(fmt_g)
        self.cb_pdf = QCheckBox('导出PDF (增强图片/表格)')
        self.cb_pdf.setChecked(True)
        self.cb_docx = QCheckBox('导出Word (增强图片/表格)')
        self.cb_docx.setChecked(True)
        fg.addWidget(self.cb_pdf)
        fg.addWidget(self.cb_docx)
        rv.addWidget(fmt_g)
        
        # 转换按钮
        self.convert_btn = QPushButton('🚀 开始导出')
        self.convert_btn.clicked.connect(self._convert)
        self.convert_btn.setEnabled(False)
        rv.addWidget(self.convert_btn)
        
        # 状态和进度
        self.conv_status = StatusIndicator(right)
        self.conv_status.setObjectName('status_bar')
        rv.addWidget(self.conv_status)
        
        self.progress = QProgressBar()
        self.progress.setVisible(False)
        rv.addWidget(self.progress)
        
        # 日志
        self.log = QTextEdit()
        self.log.setObjectName('log')
        self.log.setReadOnly(True)
        self.log.setMaximumHeight(240)
        rv.addWidget(self.log)

        # 分割器
        spl = QSplitter(Qt.Horizontal)
        spl.addWidget(left)
        spl.addWidget(right)
        spl.setStretchFactor(0, 1)  
        spl.setStretchFactor(1, 1)
        spl.setSizes([680, 420])
        spl.setHandleWidth(3)
        spl.setStyleSheet("""
            QSplitter::handle {
                background: #cbd5e1;
                border: 1px solid #94a3b8;
            }
            QSplitter::handle:horizontal {
                width: 3px;
            }
        """)
        
        main.addWidget(spl)

    def _setup_tree_columns(self):
        """设置树控件列为真正的50-50分割"""
        try:
            tree_width = self.tree.width() - 20
            col_width = tree_width // 2
            
            self.tree.setColumnWidth(0, col_width)
            self.tree.setColumnWidth(1, col_width)
            
            header = self.tree.header()
            header.setSectionResizeMode(0, header.Stretch)
            header.setSectionResizeMode(1, header.Stretch)
            header.setDefaultSectionSize(col_width)
            
        except Exception as e:
            self.logger.debug(f"设置列宽失败: {e}")

    def _apply_styles(self):
        """简洁的白色主题样式，完全无黑色"""
        style = """
        QMainWindow { 
            background: white;
            color: #374151;
        }
        
        QWidget {
            background: white;
            color: #374151;
        }
        
        QGroupBox { 
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 8px; 
            padding: 15px; 
            margin-top: 12px;
            font-size: 14px;
        }
        
        QGroupBox::title { 
            subcontrol-origin: margin; 
            left: 12px; 
            padding: 0 8px; 
            color: #1f2937; 
            font-weight: 600;
            background: white;
        }
        
        QLabel#title_label { 
            font-size: 20px; 
            font-weight: 700; 
            color: #1f2937; 
            background: white;
            padding: 10px;
            border: 1px solid #e5e7eb;
            border-radius: 6px;
        }
        
        QLabel { 
            color: #374151;
            background: white;
        }
        
        QPushButton { 
            background: #3b82f6;
            color: white; 
            border: none;
            padding: 10px 16px; 
            border-radius: 6px; 
            font-weight: 600;
            font-size: 13px;
        }
        
        QPushButton:hover { 
            background: #2563eb;
        }
        
        QPushButton:disabled { 
            background: #e5e7eb; 
            color: #9ca3af;
        }
        
        QTreeWidget { 
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 6px;
            font-size: 13px;
            color: #374151;
            outline: none;
        }
        
        QHeaderView::section {
            background: #f9fafb;
            color: #374151;
            border: none;
            border-right: 1px solid #e5e7eb;
            padding: 8px;
            font-weight: 600;
        }
        
        QTreeWidget::item {
            background: white;
            color: #374151;
            padding: 6px;
            height: 26px;
        }
        
        QTreeWidget::item:hover {
            background: #f3f4f6;
        }
        
        QTreeWidget::item:selected {
            background: #dbeafe;
            color: #1e40af;
        }
        
        QCheckBox {
            color: #374151;
            background: white;
        }
        
        QCheckBox::indicator {
            width: 16px;
            height: 16px;
            border: 1px solid #d1d5db;
            border-radius: 3px;
            background: white;
        }
        
        QCheckBox::indicator:checked {
            background: #3b82f6;
            border: 1px solid #3b82f6;
        }
        
        QTextEdit#log { 
            background: white;
            color: #374151; 
            border: 1px solid #e5e7eb;
            border-radius: 6px;
            font-family: 'Consolas', monospace;
            font-size: 11px;
            padding: 8px;
        }
        
        QProgressBar { 
            border: 1px solid #e5e7eb;
            border-radius: 6px; 
            height: 22px; 
            text-align: center;
            background: white;
            color: #374151;
        }
        
        QProgressBar::chunk { 
            background: #10b981;
            border-radius: 4px;
            margin: 1px;
        }
        """
        self.setStyleSheet(style)

    # ---- 动作 ----
    def _auto_detect(self):
        """自动检测OneNote，延迟启动避免界面卡顿"""
        if self.isVisible():
            QTimer.singleShot(200, self._refresh)
        else:
            QTimer.singleShot(1000, self._auto_detect)

    def _set_busy(self, busy: bool):
        """设置忙碌状态，禁用/启用控件"""
        self._busy=busy
        
        controls = [self.refresh_btn,self.btn_all,self.btn_none,self.btn_dir,
                   self.cb_pdf,self.cb_docx,self.convert_btn]
        
        for w in controls:
            w.setEnabled(not busy)
        
        self.tree.setEnabled(not busy)
        
        if busy:
            QApplication.setOverrideCursor(Qt.WaitCursor)
        else:
            QApplication.restoreOverrideCursor()

    def _log(self, msg: str):
        ts = QDateTime.currentDateTime().toString('hh:mm:ss')
        self.log.append(f'[{ts}] {msg}')
        self.log.verticalScrollBar().setValue(self.log.verticalScrollBar().maximum())

    def _refresh(self):
        if self._busy: return
        self._set_busy(True)
        self.refresh_status.show_loading('🔍 正在检测OneNote...')
        self.tree.clear()
        self._log('开始加载笔记本...')
        
        if hasattr(self, '_item_cache'):
            self._item_cache.clear()
        
        if self._loading_thread and self._loading_thread.isRunning():
            self._loading_thread.terminate()
            self._loading_thread.wait(100)
        
        self._loading_thread = _DetectWorker(self.onenote)
        self._loading_thread.progress.connect(self._on_detect_progress, Qt.QueuedConnection)
        self._loading_thread.done.connect(self._on_loaded, Qt.QueuedConnection)
        self._loading_thread.err.connect(self._on_load_err, Qt.QueuedConnection)
        self._loading_thread.start(QThread.HighPriority)
    
    def _on_detect_progress(self, msg: str):
        """处理检测进度"""
        self.refresh_status.show_loading(msg)
        self._log(msg)

    def _on_loaded(self, notebooks: dict):
        self.refresh_status.show_loading('📚 读取笔记本 0%')
        self._log('📚 开始读取笔记本页面...')
        
        if self._populate_thread and self._populate_thread.isRunning():
            self._populate_thread.terminate()
            self._populate_thread.wait(100)
        
        self._populate_thread = _PopulateWorker(notebooks)
        self._populate_thread.all_data.connect(self._build_tree_fast, Qt.QueuedConnection)
        self._populate_thread.progress.connect(self._on_populate_progress, Qt.QueuedConnection)
        self._populate_thread.msg.connect(self._log, Qt.QueuedConnection)
        self._populate_thread.done.connect(self._on_pop_done, Qt.QueuedConnection)
        self._populate_thread.err.connect(self._on_pop_err, Qt.QueuedConnection)
        self._populate_thread.start(QThread.HighPriority)
    
    def _on_populate_progress(self, percent: int):
        """处理构建进度"""
        if percent < 100:
            self.refresh_status.show_loading(f'📚 读取笔记本 {percent}%')

    def _on_load_err(self, msg: str):
        self.refresh_status.hide_loading(); self._set_busy(False)
        self._log(f'❌ 加载失败: {msg}')

    def _build_tree_fast(self, notebooks: dict):
        """超高速构建整个树形结构"""
        try:
            self.tree.setUpdatesEnabled(False)
            self.tree.blockSignals(True)
            self.tree.setVisible(False)
            
            self._notebooks_data = notebooks
            self._build_items = []
            self._build_index = 0
            
            for nb_id, nb_data in notebooks.items():
                nb_name = nb_data['name']
                self._build_items.append(('notebook', None, nb_id, nb_name))
                
                for sec_id, sec_data in nb_data.get('sections', {}).items():
                    sec_name = sec_data['name']
                    self._build_items.append(('section', nb_id, sec_id, sec_name))
                    
                    for page_id, page_data in sec_data.get('pages', {}).items():
                        page_name = page_data['name']
                        self._build_items.append(('page', sec_id, page_id, page_name))
            
            self._item_cache = {}
            self._build_timer = QTimer()
            self._build_timer.timeout.connect(self._build_batch)
            self._build_timer.start(1)
            
        except Exception as e:
            self._log(f'❌ 快速构建失败: {e}')
            self._finish_build()
    
    def _build_batch(self):
        """分批构建树项目"""
        try:
            batch_size = 100
            end_index = min(self._build_index + batch_size, len(self._build_items))
            
            for i in range(self._build_index, end_index):
                item_type, parent_id, item_id, item_name = self._build_items[i]
                
                if item_type == 'notebook':
                    it = QTreeWidgetItem(self.tree)
                    it.setText(0, f'📚 {item_name}')
                    it.setText(1, '笔记本')
                    it.setFlags(it.flags() | Qt.ItemIsUserCheckable)
                    it.setCheckState(0, Qt.Unchecked)
                    it.setData(0, Qt.UserRole, {'type': 'notebook', 'id': item_id, 'name': item_name})
                    it.setExpanded(True)
                    self._item_cache[item_id] = it
                    
                elif item_type == 'section':
                    parent = self._item_cache.get(parent_id)
                    if parent:
                        it = QTreeWidgetItem(parent)
                        it.setText(0, f'📁 {item_name}')
                        it.setText(1, '分区')
                        it.setFlags(it.flags() | Qt.ItemIsUserCheckable)
                        it.setCheckState(0, Qt.Unchecked)
                        it.setData(0, Qt.UserRole, {'type': 'section', 'id': item_id, 'name': item_name})
                        self._item_cache[item_id] = it
                
                elif item_type == 'page':
                    parent = self._item_cache.get(parent_id)
                    if parent:
                        it = QTreeWidgetItem(parent)
                        it.setText(0, f'📄 {item_name}')
                        it.setText(1, '页面')
                        it.setFlags(it.flags() | Qt.ItemIsUserCheckable)
                        it.setCheckState(0, Qt.Unchecked)
                        it.setData(0, Qt.UserRole, {'type': 'page', 'id': item_id, 'name': item_name})
            
            self._build_index = end_index
            
            progress = int(self._build_index * 100 / len(self._build_items))
            if progress % 10 == 0 and progress < 100:
                self.refresh_status.show_loading(f'📚 读取笔记本 {progress}%')
            
            if self._build_index >= len(self._build_items):
                self._build_timer.stop()
                self._finish_build()
                
        except Exception as e:
            self._log(f'❌ 批处理失败: {e}')
            self._build_timer.stop()
            self._finish_build()
    
    def _finish_build(self):
        """完成构建"""
        try:
            self.tree.setVisible(True)
            self.tree.blockSignals(False)
            self.tree.setUpdatesEnabled(True)
            self.refresh_status.hide_loading()
            
            if hasattr(self, '_build_items'):
                del self._build_items
            if hasattr(self, '_build_index'):
                del self._build_index
            if hasattr(self, '_notebooks_data'):
                del self._notebooks_data
                
        except Exception as e:
            self._log(f'❌ 读取时出错: {e}')

    def _find_item_by_id(self, id_: str):
        it = QTreeWidgetItemIterator(self.tree)
        while it.value():
            item = it.value()
            d=item.data(0,Qt.UserRole)
            if d and d.get('id')==id_: return item
            it+=1
        return None

    def _on_pop_done(self, nb:int, sec:int, pg:int):
        """完成界面构建"""
        self.refresh_status.hide_loading()
        self._log(f'✅ 读取完成：{nb} 笔记本，{sec} 分区，{pg} 页面')
        
        if hasattr(self, '_item_cache'):
            self._item_cache.clear()
        self._set_busy(False)

    def _on_pop_err(self, msg:str):
        """构建失败处理"""
        self.refresh_status.hide_loading()
        self._set_busy(False)
        self._log(f'❌ 构建失败: {msg}')
        if hasattr(self, '_item_cache'):
            self._item_cache.clear()

    def _on_item_changed(self, item, col):
        """处理树控件项目变化，实现级联勾选"""
        if col != 0:
            return
            
        self.tree.blockSignals(True)
        
        try:
            data = item.data(0, Qt.UserRole)
            if not data:
                return
                
            item_type = data.get('type')
            check_state = item.checkState(0)
            
            if item_type == 'notebook':
                self._cascade_check_notebook(item, check_state)
            elif item_type == 'section':
                self._cascade_check_section(item, check_state)
            elif item_type == 'page':
                self._update_parent_check_state(item)
                
        finally:
            self.tree.blockSignals(False)
            self._update_selection()
            self._update_convert()
    
    def _cascade_check_notebook(self, notebook_item, check_state):
        """级联勾选笔记本下的所有分区和页面"""
        for i in range(notebook_item.childCount()):
            section_item = notebook_item.child(i)
            section_item.setCheckState(0, check_state)
            self._cascade_check_section(section_item, check_state)
    
    def _cascade_check_section(self, section_item, check_state):
        """级联勾选分区下的所有页面"""
        for i in range(section_item.childCount()):
            page_item = section_item.child(i)
            page_item.setCheckState(0, check_state)
    
    def _update_parent_check_state(self, page_item):
        """根据子页面的勾选状态更新父分区的勾选状态"""
        section_item = page_item.parent()
        if not section_item:
            return
            
        checked_count = 0
        total_count = section_item.childCount()
        
        for i in range(total_count):
            child = section_item.child(i)
            if child.checkState(0) == Qt.Checked:
                checked_count += 1
        
        if checked_count == 0:
            section_item.setCheckState(0, Qt.Unchecked)
        elif checked_count == total_count:
            section_item.setCheckState(0, Qt.Checked)
        else:
            section_item.setCheckState(0, Qt.PartiallyChecked)
        
        self._update_notebook_check_state(section_item)
    
    def _update_notebook_check_state(self, section_item):
        """根据分区状态更新笔记本的勾选状态"""
        notebook_item = section_item.parent()
        if not notebook_item:
            return
            
        checked_count = 0
        partial_count = 0
        total_count = notebook_item.childCount()
        
        for i in range(total_count):
            child = notebook_item.child(i)
            child_state = child.checkState(0)
            if child_state == Qt.Checked:
                checked_count += 1
            elif child_state == Qt.PartiallyChecked:
                partial_count += 1
        
        if checked_count == 0 and partial_count == 0:
            notebook_item.setCheckState(0, Qt.Unchecked)
        elif checked_count == total_count:
            notebook_item.setCheckState(0, Qt.Checked)
        else:
            notebook_item.setCheckState(0, Qt.PartiallyChecked)

    def _update_selection(self):
        sel=[]; it=QTreeWidgetItemIterator(self.tree)
        while it.value():
            item=it.value(); d=item.data(0,Qt.UserRole)
            if d and d.get('type')=='page' and item.checkState(0)==Qt.Checked:
                sec=item.parent(); nb=sec.parent() if sec else None
                sel.append({'page_id': d['id'], 'page_name': d['name'], 'section_name': (sec.data(0,Qt.UserRole) or {}).get('name',''), 'notebook_name': (nb.data(0,Qt.UserRole) or {}).get('name','')})
            it+=1
        self.selected_items=sel

    def _update_convert(self):
        ok = bool(self.selected_items) and bool(self.output_dir)
        self.convert_btn.setEnabled(ok)
        self.convert_btn.setText(f'🚀 开始转换 ({len(self.selected_items)} 个页面)' if self.selected_items else '🚀 开始转换')

    def _choose_dir(self):
        d = QFileDialog.getExistingDirectory(self, '选择输出目录', self.output_dir or os.path.expanduser('~'))
        if d:
            self.output_dir=d; self.lbl_out.setText(d); self._update_convert()

    def _select_all(self):
        it=QTreeWidgetItemIterator(self.tree)
        while it.value():
            item=it.value()
            if item.flags() & Qt.ItemIsUserCheckable:
                item.setCheckState(0,Qt.Checked)
            it+=1
        self._update_selection(); self._update_convert()

    def _select_none(self):
        it=QTreeWidgetItemIterator(self.tree)
        while it.value():
            item=it.value()
            if item.flags() & Qt.ItemIsUserCheckable:
                item.setCheckState(0,Qt.Unchecked)
            it+=1
        self._update_selection(); self._update_convert()

    def _convert(self):
        if not self.selected_items or not self.output_dir:
            QMessageBox.warning(self,'提示','请选择页面和输出目录')
            return
        
        self._set_busy(True)
        self.conv_status.show_loading('🚀 正在转换...')
        self.progress.setVisible(True)
        self.progress.setValue(0)
        self.log.clear()
        
        if self._convert_thread and self._convert_thread.isRunning():
            self._convert_thread.terminate()
            self._convert_thread.wait(100)
        
        self._convert_thread = _EnhancedConvertWorker(  # 使用增强转换工作器
            self.onenote, self.parser, self.selected_items, self.output_dir,
            self.cb_pdf.isChecked(), self.cb_docx.isChecked(),
            True,  # 图片默认导出
            False  # 不支持附件
        )
        
        self._convert_thread.progress.connect(self.progress.setValue, Qt.QueuedConnection)
        self._convert_thread.msg.connect(self._log, Qt.QueuedConnection)
        self._convert_thread.done.connect(self._conv_done, Qt.QueuedConnection)
        self._convert_thread.err.connect(self._conv_err, Qt.QueuedConnection)
        self._convert_thread.start(QThread.NormalPriority)

    def _conv_done(self):
        self.progress.setValue(100); self.conv_status.hide_loading(); self._set_busy(False)
        QMessageBox.information(self,'完成','转换完成')

    def _conv_err(self, m:str):
        self.progress.setVisible(False); self.conv_status.hide_loading(); self._set_busy(False)
        QMessageBox.critical(self,'错误', m)
    
    def closeEvent(self, event):
        """关闭事件处理"""
        try:
            for thread in [getattr(self, '_loading_thread', None), 
                          getattr(self, '_populate_thread', None), 
                          getattr(self, '_convert_thread', None)]:
                if thread and thread.isRunning():
                    thread.terminate()
                    thread.wait(100)
            
            if hasattr(self, '_build_timer') and self._build_timer.isActive():
                self._build_timer.stop()
            
            if hasattr(self, 'parser'):
                self.parser.cleanup_temp_files()
                
        except Exception:
            pass
        finally:
            event.accept()
    
    def resizeEvent(self, event):
        """窗口大小改变时重新调整列宽"""
        super().resizeEvent(event)
        QTimer.singleShot(10, self._setup_tree_columns)


# ======= 线程 =======
class _DetectWorker(QThread):
    progress = pyqtSignal(str)
    done = pyqtSignal(dict)
    err = pyqtSignal(str)
    
    def __init__(self, api: OneNoteAPI):
        super().__init__()
        self.api=api
        self.setTerminationEnabled(True)
        
    def run(self):
        try:
            self.progress.emit('🔍 正在连接OneNote...')
            self.msleep(10)
            
            if not self.api.initialize():
                self.err.emit('无法连接OneNote')
                return
            
            self.progress.emit('📚 正在获取笔记本列表...')
            self.msleep(10)
            
            nbs = self.api.get_notebooks()
            if not nbs:
                self.err.emit('未发现笔记本')
                return
            
            total = sum(len(s.get('pages',{})) for nb in nbs.values() for s in nb.get('sections',{}).values())
            self.progress.emit(f'✅ 发现 {len(nbs)} 个笔记本，{total} 个页面')
            self.msleep(10)
            
            self.done.emit(nbs)
        except Exception as e:
            self.err.emit(str(e))


class _PopulateWorker(QThread):
    all_data = pyqtSignal(dict)
    progress = pyqtSignal(int)
    msg = pyqtSignal(str)
    done = pyqtSignal(int,int,int)
    err = pyqtSignal(str)
    
    def __init__(self, notebooks: dict):
        super().__init__()
        self.nbs = notebooks
        self.setTerminationEnabled(True)
        
    def run(self):
        """一次性处理所有数据，不分批"""
        try:
            nb_count = len(self.nbs)
            sec_count = 0
            pg_count = 0
            
            for nb_data in self.nbs.values():
                for sec_data in nb_data.get('sections', {}).values():
                    sec_count += 1
                    pg_count += len(sec_data.get('pages', {}))
            
            self.progress.emit(50)
            self.msleep(10)
            
            self.all_data.emit(self.nbs)
            
            self.progress.emit(100)
            self.done.emit(nb_count, sec_count, pg_count)
            
        except Exception as e:
            self.err.emit(str(e))


class _EnhancedConvertWorker(QThread):
    progress = pyqtSignal(int)
    msg = pyqtSignal(str)
    done = pyqtSignal()
    err = pyqtSignal(str)
    
    def __init__(self, api: OneNoteAPI, parser: EnhancedOneNoteContentParser, 
                 items: List[dict], out_dir: str, pdf: bool, docx: bool, 
                 images: bool, attachments: bool):
        super().__init__()
        self.api = api
        self.parser = parser
        self.items = items
        self.out = Path(out_dir)
        self.pdf = pdf
        self.docx = docx
        self.images = images
        self.attach = attachments
        
    def run(self):
        try:
            n = len(self.items)
            done = 0
            
            for it in self.items:
                try:
                    pid = it['page_id']
                    name = it['page_name']
                    nb = it['notebook_name']
                    sec = it['section_name']
                    
                    safe = lambda s: ''.join(c for c in (s or '未命名') if c.isalnum() or c in (' ','-','_','.')).strip()[:100] or '未命名'
                    d = self.out / safe(nb) / safe(sec)
                    d.mkdir(parents=True, exist_ok=True)
                    
                    # 获取页面内容，增加重试机制
                    xml = self.api.get_page_content(pid, max_retries=3)
                    if not xml: 
                        self.msg.emit(f'⚠️ 空页面: {name}')
                        continue

                    # Word导出
                    if self.docx:
                        out = d / f'{safe(name)}.docx'
                        ok = self.parser.parse_page_to_docx(xml, name, str(out), 
                                                            include_images=True,
                                                            include_attachments=False,
                                                            embed_attachments=False,
                                                            attachments_output_dir=None)
                        self.msg.emit(f'{"✅" if ok else "❌"} Word (增强): {name}')

                    # PDF导出
                    if self.pdf:
                        out = d / f'{safe(name)}.pdf'
                        ok = self.parser.parse_page_to_pdf(xml, name, str(out), 
                                                           include_images=True,
                                                           include_attachments=False,
                                                           attachments_output_dir=None)
                        self.msg.emit(f'{"✅" if ok else "❌"} PDF (增强): {name}')

                    done += 1
                    self.progress.emit(int(done / max(n, 1) * 100))
                    
                except Exception as e:
                    self.msg.emit(f"❌ 导出页面失败: {it['page_name']}，错误: {str(e)}")
                    import traceback
                    self.msg.emit(traceback.format_exc())
                    done += 1
                    self.progress.emit(int(done / max(n, 1) * 100))
                finally:
                    self.parser.cleanup_temp_files()
            
            self.done.emit()
        except Exception as e:
            self.err.emit(str(e))


def main():
    # Windows隐藏控制台：建议用 pythonw.exe 运行
    app = QApplication(sys.argv)
    
    # 设置应用程序名称
    app.setApplicationName('WATS-OneNote_Exporter')
    app.setOrganizationName('WATS')
    
    # 性能优化设置
    app.setAttribute(Qt.AA_EnableHighDpiScaling, True)  # 高DPI支持
    app.setAttribute(Qt.AA_UseHighDpiPixmaps, True)  # 高DPI图标
    app.setAttribute(Qt.AA_CompressHighFrequencyEvents, True)  # 压缩高频事件
    
    # 创建并显示主窗口
    w = ModernOneNoteGUI()
    w.show()
    
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()